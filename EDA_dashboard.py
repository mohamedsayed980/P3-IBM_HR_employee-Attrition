# =============================================================================
# ML ENGINE DASHBOARD — EDA MODULE (Part 2)
# Streamlit-based Interactive Dashboard
# Mirrors MATLAB ML_Engine.mlapp — Tabs 1 → 6
# Adds NEW: Tab 7 (Missing Values & Imputation) + Tab 8 (Multicollinearity VIF)
#
# Compatible with: ML_Engine_Step2_Outliers_Report.py  (Part 1 backend)
# Dataset tested:  kc_house_data.csv  (King County House Prices)
#
# Run with:  streamlit run ML_Engine_Dashboard.py
# =============================================================================
## path = streamlit run "E:\FINAL PROJECTS\P4_Superstore (2015-2018)\EDA_Dashboard.py"
# ================================================================#
# =============================================================================
# A — IMPORTS
# =============================================================================

# A1 — Core
import streamlit as st
import pandas as pd
import numpy as np
import os
import io

# A2 — Visualization
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import seaborn as sns

# A3 — Stats & ML
from scipy import stats
from scipy.stats import zscore
from sklearn.preprocessing import StandardScaler
from sklearn.impute import KNNImputer
from statsmodels.stats.outliers_influence import variance_inflation_factor

# A4 — Reports
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph,
    Spacer, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# =============================================================================
# B — PAGE CONFIG & GLOBAL STYLE -----> in Home.py only 
# =============================================================================
# ADD LOGO TO DASHBOARD 
import pathlib
LOGO = pathlib.Path(__file__).parent.parent / "3M_logo.png"

# =============================================================================
# C — SESSION STATE INITIALISATION
# =============================================================================

def init_state():
    defaults = {
        "df_raw"      : None,   # original loaded dataframe
        "df_clean"    : None,   # after IQR cleaning (Tab 3)
        "df_imputed"  : None,   # after imputation    (Tab 7)
        "df_work"     : None,   # working copy used across tabs
        "target_col"  : None,
        "num_cols"    : [],
        "cat_cols"    : [],
        "important_vars" : [],
        "iqr_table"   : None,   # Tab 3 outlier table
        "insights_text": "",
        "file_name"   : "",
        "corr_threshold" : 0.30,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# =============================================================================
# D — HELPER UTILITIES
# =============================================================================

def get_numeric_cols(df):
    return df.select_dtypes(include=[np.number]).columns.tolist()

def get_cat_cols(df):
    return df.select_dtypes(include=["object", "category"]).columns.tolist()

def outlier_lamp_html(pct):
    if pct < 2:
        return '<span class="badge-green">🟢 Clean (&lt;2%)</span>'
    elif pct <= 10:
        return f'<span class="badge-yellow">🟡 Moderate ({pct:.1f}%)</span>'
    else:
        return f'<span class="badge-red">🔴 Severe ({pct:.1f}%)</span>'

def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    return buf

# Set a consistent matplotlib style
plt.rcParams.update({
    "axes.spines.top"   : False,
    "axes.spines.right" : False,
    "axes.grid"         : True,
    "grid.alpha"        : 0.3,
    "font.family"       : "DejaVu Sans",
    "axes.labelcolor"   : "#1a237e",
    "axes.titlecolor"   : "#1a237e",
    "xtick.color"       : "#555",
    "ytick.color"       : "#555",
})

BLUE   = "#1565c0"
ORANGE = "#e65100"
GREEN  = "#2e7d32"
RED    = "#c62828"
TEAL   = "#00695c"

# =============================================================================
# E — HEADER
# =============================================================================

st.markdown("""
<div class="main-header">
    <h1>🔬 HR Attrition — EDA Dashboard</h1>
    <p>Exploratory Data Analysis · Attrition Drivers · Salary Patterns · HR Insights</p>
</div>
""", unsafe_allow_html=True)

# =============================================================================
# F — FILE LOADER (Sidebar-free: shown above tabs)
# =============================================================================

with st.sidebar:
    st.image(str(LOGO), width=70)
    st.markdown("---")
    

with st.container():
    col_load, col_target, col_thresh, col_info = st.columns([3, 2, 2, 3])

    with col_load:
        uploaded = st.file_uploader(
            "📂 Load Dataset (.csv)", type=["csv"],
            key="file_uploader", label_visibility="collapsed",
            help="Upload your CSV dataset"
        )
        if uploaded:
            try:
                df = pd.read_csv(uploaded, sep=None, engine="python")
                st.session_state.df_raw  = df.copy()
                st.session_state.df_work = df.copy()
                st.session_state.file_name = uploaded.name
                
                st.session_state.num_cols = get_numeric_cols(df)
                st.session_state.cat_cols = get_cat_cols(df)
                
                # Safety: if get_numeric_cols filtered too aggressively, fallback to all numeric  # add that safety
                if len(st.session_state.num_cols) == 0:
                    st.session_state.num_cols = df.select_dtypes(include="number").columns.tolist()
                    if len(st.session_state.cat_cols) == 0:
                        st.session_state.cat_cols = df.select_dtypes(include="object").columns.tolist()
                
                st.success(f"✅  Loaded **{uploaded.name}** — {df.shape[0]:,} rows × {df.shape[1]} columns")
            except Exception as e:
                st.error(f"Error loading file: {e}")

    with col_target:
        if st.session_state.df_raw is not None:
            cols = st.session_state.df_raw.columns.tolist()
            
            #default_idx = cols.index("price") if "price" in cols else 0                    # old repo1# xx
            
            default_idx = cols.index("Attrition_flag") if "Attrition_flag" in cols else 0
            
            target = st.selectbox("🎯 Target Variable", cols, index=default_idx)
            st.session_state.target_col = target

    with col_thresh:
        thresh = st.slider(
            "Correlation Threshold",
            0.10, 0.90,
            float(st.session_state.corr_threshold),
            0.05
        )
        st.session_state.corr_threshold = thresh

    with col_info:
        if st.session_state.df_raw is not None:
            df = st.session_state.df_raw
            st.markdown(f"""
            <div style="background:white;border-radius:8px;padding:10px 14px;
                        box-shadow:0 2px 6px rgba(0,0,0,.08);font-size:0.82rem;line-height:1.8;">
                📊 <b>Shape:</b> {df.shape[0]:,} × {df.shape[1]}<br>
                🔢 <b>Numeric:</b> {len(st.session_state.num_cols)} &nbsp;|&nbsp;
                🔤 <b>Categorical:</b> {len(st.session_state.cat_cols)}<br>
                ❓ <b>Missing:</b> {df.isnull().sum().sum():,} cells
            </div>
            """, unsafe_allow_html=True)
        else:
            st.info("⬆️ Please upload a CSV file to begin.")
            
     
st.markdown("---")

# =============================================================================
# G — TABS
# =============================================================================

tab_labels = [
    "📊 Tab 1 · Data & Correlation",
    "📈 Tab 2 · Variables Analysis",
    "🧹 Tab 3 · IQR Cleaning",
    "🔍 Tab 4 · Outliers Lab",
    "📋 Tab 5 · Dashboard Summary",
    "🩹 Tab 6 · Missing Values",
    "🔗 Tab 7 · Multicollinearity",
    "💡 Tab 8 · Insights",
    "📦 Tab 9 · HR KPI Dashboard",
    "🔍 Tab 10 · Attrition Deep Dive",
    "💰 Tab 11 · Salary Analysis",
    "🧪 Tab 12 · Statistical Tests",
]

tabs = st.tabs(tab_labels)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — DATA & CORRELATION
# ─────────────────────────────────────────────────────────────────────────────

with tabs[0]:

    st.markdown('<div class="section-title">📊 Data Overview & Correlation Analysis</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
        
    else:
        
        df  = st.session_state.df_work
        tgt = st.session_state.target_col
       # Always derive fresh from df — never trust session state for columns
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat
        

        # ── Row 1: Key Metrics
        m1, m2, m3, m4, m5 = st.columns(5)
        for col_w, label, val in zip(
            [m1, m2, m3, m4, m5],
            ["Rows", "Columns", "Numeric", "Missing Cells", "Duplicate Rows"],
            [f"{df.shape[0]:,}", f"{df.shape[1]}",
             f"{len(num)}",
             f"{df.isnull().sum().sum():,}",
             f"{df.duplicated().sum():,}"]
        ):
            col_w.markdown(f"""
            <div class="metric-card"><h4>{label}</h4><p>{val}</p></div>
            """, unsafe_allow_html=True)

        st.markdown("")

        col_left, col_right = st.columns([1, 1])

        # ── Left: Descriptive Statistics
        with col_left:
            st.markdown('<div class="section-title">Descriptive Statistics</div>', unsafe_allow_html=True)
            # Safety re-derive in case df_work was reset between reruns
            num = df.select_dtypes(include="number").columns.tolist()
            if len(num) == 0:
                st.warning(f"⚠️ No numeric columns found. df shape: {df.shape}, dtypes: {df.dtypes.value_counts().to_dict()}")
                st.stop()
            desc = df[num].describe().T.round(3)
            desc.index.name = "Variable"
            st.dataframe(desc, use_container_width=True, height=320)

        # ── Right: Top Correlations with Target
        with col_right:
            st.markdown(f'<div class="section-title">Top Correlations with Target: <i>{tgt}</i></div>', unsafe_allow_html=True)

            if tgt in df.columns:
                corr_series = df[num].corr()[tgt].drop(tgt, errors="ignore")
                corr_df = (
                    corr_series.abs()
                    .sort_values(ascending=False)
                    .reset_index()
                )
                corr_df.columns = ["Variable", "Abs Correlation"]
                corr_df["Correlation"] = corr_series.reindex(corr_df["Variable"]).values.round(4)

                # Color strong correlations
                def color_corr(val):
                    abs_val = abs(val)
                    if abs_val >= 0.70: return "background-color:#c8e6c9; color:#1b5e20;"
                    elif abs_val >= 0.50: return "background-color:#fff9c4; color:#e65100;"
                    else: return ""

                styled = corr_df.style.applymap(color_corr, subset=["Correlation"])
                st.dataframe(styled, use_container_width=True, height=320)

                # Update important_vars
                threshold = st.session_state.corr_threshold
                st.session_state.important_vars = corr_df[
                    corr_df["Abs Correlation"] >= threshold
                ]["Variable"].tolist()

        st.markdown("---")

        # ── Buttons Row
        btn1, btn2, btn3, _ = st.columns([1.5, 1.5, 1.5, 5])

        with btn1:
            show_heatmap = st.button("🌡️ Show Correlation Heatmap", key="btn_heatmap1")
        with btn2:
            show_scatter = st.button("📉 Show Scatter Plots", key="btn_scatter1")
        with btn3:
            show_data    = st.button("👁️ Preview Raw Data", key="btn_data1")

        if show_heatmap:
            st.markdown('<div class="section-title">Correlation Heatmap</div>', unsafe_allow_html=True)
            fig, ax = plt.subplots(figsize=(12, 9))
            plot_cols = num[:18]  # max 18 for readability
            corr_mat  = df[plot_cols].corr()
            mask = np.triu(np.ones_like(corr_mat, dtype=bool))
            sns.heatmap(
                corr_mat, mask=mask, ax=ax, cmap="RdYlGn",
                annot=True, fmt=".2f", linewidths=0.5, linecolor="#e0e0e0",
                vmin=-1, vmax=1, annot_kws={"size": 7.5},
                cbar_kws={"shrink": 0.7}
            )
            ax.set_title("Correlation Matrix (lower triangle)", fontsize=14, pad=14, weight="bold")
            plt.xticks(rotation=45, ha="right", fontsize=9)
            plt.yticks(fontsize=9)
            st.pyplot(fig, use_container_width=True)
            plt.close()

        if show_scatter:
            imp_vars = st.session_state.important_vars[:8]
            if imp_vars and tgt in df.columns:
                st.markdown(f'<div class="section-title">Scatter Plots — Top Variables vs {tgt}</div>', unsafe_allow_html=True)
                n = len(imp_vars)
                ncols = 4
                nrows = (n + ncols - 1) // ncols
                fig, axes = plt.subplots(nrows, ncols, figsize=(16, 4 * nrows))
                axes = axes.flatten() if n > 1 else [axes]
                for i, var in enumerate(imp_vars):
                    ax = axes[i]
                    ax.scatter(df[var], df[tgt], alpha=0.3, s=12, color=BLUE)
                    # Regression line
                    x_c = df[var].dropna()
                    y_c = df[tgt].dropna()
                    common = df[[var, tgt]].dropna()
                    if len(common) > 10:
                        m, b = np.polyfit(common[var], common[tgt], 1)
                        x_line = np.linspace(common[var].min(), common[var].max(), 100)
                        ax.plot(x_line, m * x_line + b, color=RED, linewidth=1.8)
                    corr_val = df[[var, tgt]].corr().iloc[0, 1]
                    ax.set_title(f"{var} vs {tgt}\n(r = {corr_val:.3f})", fontsize=9, weight="bold")
                    ax.set_xlabel(var, fontsize=8)
                    ax.set_ylabel(tgt, fontsize=8)
                for j in range(i + 1, len(axes)):
                    axes[j].set_visible(False)
                plt.tight_layout(pad=2)
                st.pyplot(fig, use_container_width=True)
                plt.close()
            else:
                st.info("No important variables found above the correlation threshold.")

        if show_data:
            st.markdown('<div class="section-title">Raw Data Preview (first 100 rows)</div>', unsafe_allow_html=True)
            st.dataframe(df.head(100), use_container_width=True, height=400)

        # Export
        st.markdown("---")
        if st.button("📥 Export Descriptive Statistics (CSV)", key="exp_tab1"):
            csv_data = df[num].describe().T.round(4).to_csv()
            st.download_button(
                "⬇️ Download CSV", csv_data,
                file_name="descriptive_statistics.csv",
                mime="text/csv"
            )


# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — VARIABLES ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

with tabs[1]:

    st.markdown('<div class="section-title">📈 Variables Analysis</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        tgt = st.session_state.target_col
        
        num = df.select_dtypes(include="number").columns.tolist()  # updated 

        col_controls, col_plot = st.columns([1, 2.5])

        with col_controls:
            st.markdown("**Select Variable**")
            imp_vars = st.session_state.important_vars if st.session_state.important_vars else num
            sel_var  = st.selectbox("Variable", imp_vars, key="tab2_var")

            plot_type = st.radio(
                "Plot Type",
                ["Scatter vs Target", "Histogram", "Boxplot"],
                key="tab2_plottype"
            )

            if sel_var:
                col_stats = df[sel_var].describe().round(3)
                st.markdown("**Quick Stats**")
                stats_html = "".join([
                    f"<div style='display:flex;justify-content:space-between;"
                    f"padding:4px 0;border-bottom:1px solid #eee;font-size:0.82rem;'>"
                    f"<span style='color:#555;'>{k}</span>"
                    f"<span style='font-weight:600;color:#1a237e;'>{v:.3f}</span></div>"
                    for k, v in col_stats.items()
                ])
                skew_val = df[sel_var].skew()
                kurt_val = df[sel_var].kurt()
                stats_html += (
                    f"<div style='display:flex;justify-content:space-between;"
                    f"padding:4px 0;border-bottom:1px solid #eee;font-size:0.82rem;'>"
                    f"<span style='color:#555;'>Skewness</span>"
                    f"<span style='font-weight:600;color:#1a237e;'>{skew_val:.3f}</span></div>"
                    f"<div style='display:flex;justify-content:space-between;"
                    f"padding:4px 0;font-size:0.82rem;'>"
                    f"<span style='color:#555;'>Kurtosis</span>"
                    f"<span style='font-weight:600;color:#1a237e;'>{kurt_val:.3f}</span></div>"
                )
                st.markdown(
                    f"<div style='background:white;border-radius:8px;"
                    f"padding:12px 14px;box-shadow:0 2px 6px rgba(0,0,0,.08);'>"
                    f"{stats_html}</div>",
                    unsafe_allow_html=True
                )

        with col_plot:
            if sel_var:
                fig, ax = plt.subplots(figsize=(9, 5.5))
                data_clean = df[[sel_var]].dropna()

                if plot_type == "Scatter vs Target" and tgt in df.columns:
                    xy = df[[sel_var, tgt]].dropna()
                    ax.scatter(xy[sel_var], xy[tgt], alpha=0.35, s=14, color=BLUE, label="Data")
                    if len(xy) > 10:
                        m, b = np.polyfit(xy[sel_var], xy[tgt], 1)
                        x_l = np.linspace(xy[sel_var].min(), xy[sel_var].max(), 200)
                        ax.plot(x_l, m * x_l + b, color=RED, linewidth=2, label="Trend line")
                    corr_v = df[[sel_var, tgt]].corr().iloc[0, 1]
                    ax.set_title(f"{tgt} vs {sel_var}  (r = {corr_v:.3f})", weight="bold", fontsize=12)
                    ax.set_xlabel(sel_var); ax.set_ylabel(tgt)
                    ax.legend()

                elif plot_type == "Histogram":
                    d = data_clean[sel_var]
                    ax.hist(d, bins=40, color=BLUE, alpha=0.75, edgecolor="white", linewidth=0.5)
                    ax.axvline(d.mean(),  color=RED,    linestyle="--", linewidth=1.8, label=f"Mean={d.mean():.2f}")
                    ax.axvline(d.median(), color=GREEN, linestyle="--", linewidth=1.8, label=f"Median={d.median():.2f}")
                    ax.set_title(f"Distribution of {sel_var}", weight="bold", fontsize=12)
                    ax.set_xlabel(sel_var); ax.set_ylabel("Frequency")
                    ax.legend()

                elif plot_type == "Boxplot":
                    bp = ax.boxplot(
                        data_clean[sel_var].values,
                        patch_artist=True,
                        notch=False,
                        vert=True,
                        widths=0.5,
                        boxprops=dict(facecolor="#bbdefb", color=BLUE),
                        medianprops=dict(color=RED, linewidth=2.5),
                        whiskerprops=dict(color=BLUE),
                        capprops=dict(color=BLUE),
                        flierprops=dict(marker="o", color=ORANGE, alpha=0.5, markersize=4)
                    )
                    ax.set_title(f"Boxplot of {sel_var}", weight="bold", fontsize=12)
                    ax.set_ylabel(sel_var)
                    ax.set_xticks([])

                plt.tight_layout()
                st.pyplot(fig, use_container_width=True)
                plt.close()

                # Normality hint
                if plot_type == "Histogram":
                    sk = abs(df[sel_var].skew())
                    if sk < 0.5:
                        msg = f"✅ Distribution is approximately symmetric (skewness={sk:.2f}). Good for regression."
                        st.markdown(f'<div class="insight-box">{msg}</div>', unsafe_allow_html=True)
                    else:
                        msg = f"⚠️ Distribution is skewed (skewness={sk:.2f}). Consider log transformation."
                        st.markdown(f'<div class="warning-box">{msg}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — IQR CLEANING
# ─────────────────────────────────────────────────────────────────────────────

with tabs[2]:

    st.markdown('<div class="section-title">🧹 IQR Cleaner — Outlier Detection & Treatment</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work.copy()
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        col_ctrl, col_plots = st.columns([1, 2.5])

        with col_ctrl:
            sel_iqr_var = st.selectbox("Variable to Inspect", num, key="iqr_var")

            iqr_action  = st.radio(
                "Outlier Action",
                ["Cap (Winsorise)", "Remove", "Keep"],
                key="iqr_action"
            )
            multiplier  = st.slider("IQR Multiplier", 1.0, 3.0, 1.5, 0.1, key="iqr_mult")

            if st.button("🔬 Compute All Outliers", key="btn_compute_iqr"):
                rows = []
                for col in num:
                    Q1  = df[col].quantile(0.25)
                    Q3  = df[col].quantile(0.75)
                    IQR = Q3 - Q1
                    lo  = Q1 - multiplier * IQR
                    hi  = Q3 + multiplier * IQR
                    mask = (df[col] < lo) | (df[col] > hi)
                    rows.append({
                        "Variable"     : col,
                        "Outlier Count": int(mask.sum()),
                        "Outlier %"    : round(mask.mean() * 100, 2),
                        "Lower Bound"  : round(lo, 3),
                        "Upper Bound"  : round(hi, 3),
                        "Action"       : "Keep"
                    })
                st.session_state.iqr_table = pd.DataFrame(rows)

            st.markdown("")
            if st.button("✅ Apply Cleaning & Save", key="btn_apply_iqr", type="primary"):
                if st.session_state.iqr_table is not None:
                    df_new = st.session_state.df_work.copy()
                    iqr_df = st.session_state.iqr_table
                    for _, row in iqr_df.iterrows():
                        col = row["Variable"]
                        act = row["Action"]
                        lo  = row["Lower Bound"]
                        hi  = row["Upper Bound"]
                        if act == "Cap (Winsorise)":
                            df_new[col] = df_new[col].clip(lower=lo, upper=hi)
                        elif act == "Remove":
                            df_new = df_new[
                                (df_new[col] >= lo) & (df_new[col] <= hi)
                            ]
                    st.session_state.df_clean = df_new.copy()
                    st.session_state.df_work  = df_new.copy()
                    st.success(f"✅ Cleaning applied. Dataset now has {df_new.shape[0]:,} rows.")
                else:
                    st.warning("Run 'Compute All Outliers' first.")

        with col_plots:
            # Show IQR table if available
            if st.session_state.iqr_table is not None:
                iqr_df = st.session_state.iqr_table

                # Summary stats
                total_outliers = iqr_df["Outlier Count"].sum()
                worst_var      = iqr_df.loc[iqr_df["Outlier Count"].idxmax(), "Variable"]
                worst_pct      = iqr_df["Outlier %"].max()

                mc1, mc2, mc3 = st.columns(3)
                mc1.markdown(f'<div class="metric-card"><h4>Total Outliers</h4><p>{total_outliers:,}</p></div>', unsafe_allow_html=True)
                mc2.markdown(f'<div class="metric-card"><h4>Most Affected</h4><p>{worst_var}</p></div>', unsafe_allow_html=True)
                mc3.markdown(f'<div class="metric-card"><h4>Max Outlier %</h4><p>{worst_pct:.1f}%</p></div>', unsafe_allow_html=True)

                # Styled table
                def style_outlier_pct(val):
                    if val < 2:   return "background-color:#e8f5e9; color:#2e7d32;"
                    elif val <= 10: return "background-color:#fff8e1; color:#e65100;"
                    else:          return "background-color:#ffebee; color:#c62828;"

                styled_iqr = iqr_df.style.applymap(style_outlier_pct, subset=["Outlier %"])
                st.dataframe(styled_iqr, use_container_width=True, height=230)

                st.markdown("")

            # Before / After boxplots for selected variable
            if sel_iqr_var:
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5), sharey=False)

                raw_data = st.session_state.df_raw[sel_iqr_var].dropna()

                # Left: Raw
                ax1.boxplot(
                    raw_data.values,
                    patch_artist=True,
                    boxprops=dict(facecolor="#ffcdd2", color=RED),
                    medianprops=dict(color=RED, linewidth=2.5),
                    flierprops=dict(marker="o", color=RED, alpha=0.4, markersize=5)
                )
                ax1.set_title("⚠️ Before Cleaning", weight="bold", color=RED)
                ax1.set_ylabel(sel_iqr_var)
                ax1.set_xticks([])

                # Right: After (if cleaned)
                if st.session_state.df_clean is not None and sel_iqr_var in st.session_state.df_clean:
                    clean_data = st.session_state.df_clean[sel_iqr_var].dropna()
                else:
                    # Show capped preview
                    Q1, Q3 = raw_data.quantile(0.25), raw_data.quantile(0.75)
                    IQR = Q3 - Q1
                    lo, hi = Q1 - multiplier * IQR, Q3 + multiplier * IQR
                    clean_data = raw_data.clip(lower=lo, upper=hi)

                ax2.boxplot(
                    clean_data.values,
                    patch_artist=True,
                    boxprops=dict(facecolor="#c8e6c9", color=GREEN),
                    medianprops=dict(color=GREEN, linewidth=2.5),
                    flierprops=dict(marker="o", color=TEAL, alpha=0.4, markersize=5)
                )
                ax2.set_title("✅ After Cleaning", weight="bold", color=GREEN)
                ax2.set_ylabel(sel_iqr_var)
                ax2.set_xticks([])

                plt.suptitle(
                    f"Outlier View: {sel_iqr_var}  "
                    f"(removed {len(raw_data) - len(clean_data):,} rows)",
                    fontsize=11, weight="bold", y=1.02
                )
                plt.tight_layout()
                st.pyplot(fig, use_container_width=True)
                plt.close()

        # Export cleaned data
        if st.session_state.df_clean is not None:
            st.markdown("---")
            csv_clean = st.session_state.df_clean.to_csv(index=False)
            st.download_button(
                "📥 Download Cleaned Dataset (CSV)",
                csv_clean,
                file_name="data_cleaned.csv",
                mime="text/csv",
                key="dl_clean"
            )


# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — OUTLIERS DETECTION LAB
# ─────────────────────────────────────────────────────────────────────────────

with tabs[3]:

    st.markdown('<div class="section-title">🔍 Outliers Detection Lab</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        col_lab_ctrl, col_lab_vis = st.columns([1, 2.5])

        with col_lab_ctrl:
            lab_var    = st.selectbox("Variable", num, key="lab_var")
            lab_method = st.radio("Detection Method", ["IQR (Q1-Q3)", "Z-Score"], key="lab_method")
            lab_thresh = st.number_input(
                "Threshold (IQR multiplier or Z-Score)",
                min_value=0.5, max_value=5.0, value=1.5, step=0.1, key="lab_thresh"
            )

            run_analysis = st.button("🔬 Analyze", key="btn_lab_analyze", type="primary")

        with col_lab_vis:
            if run_analysis and lab_var:
                data = df[lab_var].dropna()
                n    = len(data)

                # Detect outliers
                if lab_method == "IQR (Q1-Q3)":
                    Q1, Q3 = data.quantile(0.25), data.quantile(0.75)
                    IQR    = Q3 - Q1
                    lo_b   = Q1 - lab_thresh * IQR
                    hi_b   = Q3 + lab_thresh * IQR
                    is_out = (data < lo_b) | (data > hi_b)
                else:
                    zs     = np.abs(zscore(data))
                    is_out = zs > lab_thresh
                    lo_b   = data.mean() - lab_thresh * data.std()
                    hi_b   = data.mean() + lab_thresh * data.std()

                out_count = is_out.sum()
                out_pct   = out_count / n * 100

                # Status card
                badge_html = outlier_lamp_html(out_pct)
                sc1, sc2, sc3 = st.columns(3)
                sc1.markdown(f'<div class="metric-card"><h4>Outliers Found</h4><p>{out_count:,}</p></div>', unsafe_allow_html=True)
                sc2.markdown(f'<div class="metric-card"><h4>Percentage</h4><p>{out_pct:.2f}%</p></div>', unsafe_allow_html=True)
                sc3.markdown(f'<div class="metric-card"><h4>Data Quality</h4><br>{badge_html}</div>', unsafe_allow_html=True)

                # Two plots: Distribution + Scatter highlight
                fig, (ax_top, ax_bot) = plt.subplots(2, 1, figsize=(10, 8))

                # ─ Distribution plot
                ax_top.hist(data[~is_out], bins=40, color=BLUE, alpha=0.7, label="Normal", density=True)
                ax_top.hist(data[is_out],  bins=20, color=RED,  alpha=0.8, label="Outliers", density=True)
                ax_top.axvline(lo_b, color=ORANGE, linestyle="--", linewidth=1.8, label=f"Lower limit = {lo_b:.1f}")
                ax_top.axvline(hi_b, color=ORANGE, linestyle="--", linewidth=1.8, label=f"Upper limit = {hi_b:.1f}")
                ax_top.axvline(data.mean(), color=GREEN, linestyle="-", linewidth=1.5, label=f"μ = {data.mean():.1f}")
                ax_top.set_title(
                    f"Distribution of {lab_var}   (σ = {data.std():.2f}  |  method: {lab_method})",
                    fontsize=11, weight="bold"
                )
                ax_top.set_xlabel(lab_var); ax_top.set_ylabel("Density")
                ax_top.legend(fontsize=8)

                # ─ Scatter outlier highlight
                idx_arr = np.arange(n)
                ax_bot.scatter(idx_arr[~is_out], data[~is_out], s=8,  color=BLUE, alpha=0.4, label="Normal")
                ax_bot.scatter(idx_arr[is_out],  data[is_out],  s=18, color=RED,  alpha=0.75, label="Outlier", zorder=5)
                ax_bot.axhline(lo_b, color=ORANGE, linestyle="--", linewidth=1.5)
                ax_bot.axhline(hi_b, color=ORANGE, linestyle="--", linewidth=1.5)
                ax_bot.set_title("Outlier Highlight — Scatter View", fontsize=11, weight="bold")
                ax_bot.set_xlabel("Row Index"); ax_bot.set_ylabel(lab_var)
                ax_bot.legend(fontsize=8)

                plt.tight_layout(pad=2)
                st.pyplot(fig, use_container_width=True)
                plt.close()

                # Outlier values table
                out_vals = data[is_out].sort_values(ascending=False).head(50)
                if len(out_vals) > 0:
                    st.markdown('<div class="section-title">Top Outlier Values</div>', unsafe_allow_html=True)
                    out_df = out_vals.reset_index()
                    out_df.columns = ["Row Index", f"{lab_var} (outlier value)"]
                    st.dataframe(out_df, use_container_width=True, height=220)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — DASHBOARD SUMMARY
# ─────────────────────────────────────────────────────────────────────────────

with tabs[4]:

    st.markdown('<div class="section-title">📋 Summary & Correlation Insights</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        tgt = st.session_state.target_col
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        btn_s1, btn_s2, _ = st.columns([2, 2, 6])
        gen_summary = btn_s1.button("📊 Generate Summary Table", key="btn_gen_summary")
        gen_heat    = btn_s2.button("🌡️ Generate Heatmap",       key="btn_gen_heat")

        if gen_summary or "summary_table" in st.session_state:
            if gen_summary:
                if tgt in df.columns:
                    corr_vals = df[num].corr()[tgt]
                    desc      = df[num].describe().T

                    summary_data = {
                        "Variable"          : num,
                        "Mean"              : [round(desc.loc[c, "mean"], 3) if c in desc.index else None for c in num],
                        "Std Dev"           : [round(desc.loc[c, "std"],  3) if c in desc.index else None for c in num],
                        "Min"               : [round(desc.loc[c, "min"],  3) if c in desc.index else None for c in num],
                        "Max"               : [round(desc.loc[c, "max"],  3) if c in desc.index else None for c in num],
                        "Corr with Target"  : [round(corr_vals.get(c, 0), 4) for c in num],
                        "Skewness"          : [round(df[c].skew(), 3) for c in num],
                        "Missing %"         : [round(df[c].isnull().mean() * 100, 2) for c in num],
                    }
                    st.session_state["summary_table"] = pd.DataFrame(summary_data)

            if "summary_table" in st.session_state:
                sum_df = st.session_state["summary_table"]

                col_imp, col_table = st.columns([1, 3])
                with col_imp:
                    st.markdown("**Important Variables**")
                    imp_html = "".join([
                        f"<div style='padding:5px 10px;margin-bottom:4px;background:#e3f2fd;"
                        f"border-radius:5px;font-size:0.83rem;color:#0d47a1;font-weight:600;'>{v}</div>"
                        for v in st.session_state.important_vars[:12]
                    ])
                    st.markdown(imp_html, unsafe_allow_html=True)

                with col_table:
                    def color_corr_summary(val):
                        try:
                            abs_v = abs(float(val))
                            if abs_v >= 0.7: return "background:#c8e6c9;color:#1b5e20;font-weight:bold;"
                            elif abs_v >= 0.5: return "background:#fff9c4;color:#e65100;"
                            return ""
                        except: return ""

                    styled_sum = sum_df.style.applymap(color_corr_summary, subset=["Corr with Target"])
                    st.dataframe(styled_sum, use_container_width=True, height=380)

                # Export summary
                st.markdown("---")
                ec1, ec2, _ = st.columns([2, 2, 6])
                with ec1:
                    csv_s = sum_df.to_csv(index=False)
                    st.download_button("📥 Export CSV", csv_s, "summary_table.csv", "text/csv", key="dl_sum_csv")

        if gen_heat:
            st.markdown("---")
            st.markdown('<div class="section-title">Correlation Heatmap</div>', unsafe_allow_html=True)
            plot_cols = [c for c in num if c in df.columns][:18]
            corr_mat  = df[plot_cols].corr()
            fig, ax   = plt.subplots(figsize=(13, 10))
            mask      = np.triu(np.ones_like(corr_mat, dtype=bool))
            sns.heatmap(
                corr_mat, mask=mask, ax=ax,
                cmap="RdYlGn", annot=True, fmt=".2f",
                linewidths=0.5, linecolor="#e0e0e0",
                vmin=-1, vmax=1, annot_kws={"size": 8},
                cbar_kws={"shrink": 0.7, "label": "Pearson r"}
            )
            ax.set_title(
                f"Correlation Heatmap — {len(plot_cols)} Variables\n"
                f"(Target: {tgt}  |  Threshold: {st.session_state.corr_threshold})",
                fontsize=13, pad=16, weight="bold"
            )
            plt.xticks(rotation=40, ha="right", fontsize=9)
            plt.yticks(fontsize=9)
            st.pyplot(fig, use_container_width=True)
            plt.close()


# ─────────────────────────────────────────────────────────────────────────────
# TAB 8 — INSIGHTS & RECOMMENDATIONS
# ─────────────────────────────────────────────────────────────────────────────

with tabs[7]:

    st.markdown('<div class="section-title">💡 Insights & Recommendations Overview</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        tgt = st.session_state.target_col
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        if st.button("⚡ Generate Recommendations", key="btn_gen_rec", type="primary"):

            lines = []
            lines.append("=" * 60)
            lines.append("  INSIGHTS & RECOMMENDATIONS REPORT")
            lines.append(f"  Dataset : {st.session_state.file_name}")
            lines.append(f"  Target  : {tgt}")
            lines.append(f"  Shape   : {df.shape[0]:,} rows × {df.shape[1]} columns")
            lines.append("=" * 60)
            lines.append("")

            # 1 — Correlation findings
            lines.append("1. CORRELATION FINDINGS")
            lines.append("─" * 40)
            if tgt in df.columns:
                corr_s = df[num].corr()[tgt].drop(tgt, errors="ignore")
                top5   = corr_s.abs().sort_values(ascending=False).head(5)
                lines.append(f"   Variables most correlated with '{tgt}':")
                for var, cval in top5.items():
                    real_c = corr_s[var]
                    sign   = "positive" if real_c > 0 else "negative"
                    lines.append(f"   • {var:<20} r = {real_c:+.4f}  ({sign})")
                best = top5.index[0]
                lines.append(f"\n   KEY INSIGHT: '{best}' shows the strongest")
                lines.append(f"   correlation (r = {corr_s[best]:+.4f}) with {tgt}.")
                lines.append("")

            # 2 — Data quality
            lines.append("2. DATA QUALITY ASSESSMENT")
            lines.append("─" * 40)
            missing_total = df.isnull().sum().sum()
            lines.append(f"   Total missing cells : {missing_total:,}")
            if missing_total > 0:
                miss_cols = df.isnull().sum()
                miss_cols = miss_cols[miss_cols > 0].sort_values(ascending=False)
                for col, cnt in miss_cols.items():
                    pct = cnt / len(df) * 100
                    lines.append(f"   • {col:<20} {cnt:,} missing ({pct:.1f}%)")
                lines.append("   RECOMMENDATION: Impute missing values")
                lines.append("   before training models (see Tab 6).")
            else:
                lines.append("   ✅ No missing values detected.")
            lines.append("")

            # 3 — Outlier summary
            lines.append("3. OUTLIER SUMMARY (IQR method, 1.5×)")
            lines.append("─" * 40)
            for col in num[:10]:
                Q1, Q3 = df[col].quantile(0.25), df[col].quantile(0.75)
                IQR    = Q3 - Q1
                lo, hi = Q1 - 1.5 * IQR, Q3 + 1.5 * IQR
                out_n  = ((df[col] < lo) | (df[col] > hi)).sum()
                out_p  = out_n / len(df) * 100
                status = "SEVERE" if out_p > 10 else ("MODERATE" if out_p > 2 else "CLEAN")
                lines.append(f"   • {col:<20} {out_n:>5,} outliers  ({out_p:.1f}%)  [{status}]")
            lines.append("")

            # 4 — Skewness check
            lines.append("4. SKEWNESS ASSESSMENT")
            lines.append("─" * 40)
            high_skew = []
            for col in num:
                sk = abs(df[col].skew())
                if sk > 1:
                    high_skew.append((col, df[col].skew()))
            if high_skew:
                lines.append("   Highly skewed variables (|skew| > 1):")
                for col, sk in sorted(high_skew, key=lambda x: abs(x[1]), reverse=True)[:8]:
                    lines.append(f"   • {col:<22} skewness = {sk:.3f}")
                lines.append("   RECOMMENDATION: Apply log/sqrt transformation")
                lines.append("   to these variables before modelling.")
            else:
                lines.append("   ✅ No severely skewed variables detected.")
            lines.append("")

            # 5 — Recommendation summary
            lines.append("5. MODELLING RECOMMENDATIONS")
            lines.append("─" * 40)
            lines.append(f"   • Use the {len(st.session_state.important_vars)} variables with")
            lines.append(f"     corr ≥ {st.session_state.corr_threshold} as features.")
            lines.append("   • Run VIF check (Tab 7) to detect")
            lines.append("     multicollinearity before Regression.")
            lines.append("   • Apply IQR cleaning (Tab 3) to remove")
            lines.append("     or cap extreme outliers.")
            lines.append("   • Handle any missing values (Tab 6)")
            lines.append("     before splitting train/test.")
            lines.append("")
            lines.append("=" * 60)
            lines.append("  Report generated by ML Engine EDA Dashboard")
            lines.append("=" * 60)

            st.session_state.insights_text = "\n".join(lines)

        # Display text
        if st.session_state.insights_text:
            st.text_area(
                "Insights & Recommendations",
                value=st.session_state.insights_text,
                height=480,
                key="insights_area"
            )
#----------------------------------------------------
        # ── Stage 2 results bridge ────────────────────────────────────────
        stage2 = st.session_state.get("stage2_insights", "")
        if stage2:
            st.markdown(
                '<div class="info-box">📊 Stage 2 ML Results are available below.</div>',
                unsafe_allow_html=True
            )
            st.text_area(
                "📊 Stage 2 ML Results",
                value=stage2,
                height=200,
                key="stage2_area"
            )
#===============================================================        
            #st.markdown("---")
            exp_col1, exp_col2, _ = st.columns([2, 2, 6])
#------- --------------------------------


            #st.markdown("---")
            exp_col1, exp_col2, _ = st.columns([2, 2, 6])

            with exp_col1:
                st.download_button(
                    "📥 Export as TXT",
                    data=st.session_state.insights_text,
                    file_name="insights_recommendations.txt",
                    mime="text/plain",
                    key="dl_insights_txt"
                )

            with exp_col2:
                # Word export
                if st.button("📄 Export as Word (.docx)", key="btn_exp_word_insights"):
                    doc = Document()
                    doc.add_heading("Insights & Recommendations Report", 0)
                    for line in st.session_state.insights_text.split("\n"):
                        p = doc.add_paragraph(line)
                        p.style.font.size = Pt(10)
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button(
                        "⬇️ Download .docx", buf,
                        file_name="insights_recommendations.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_word_insights"
                    )


# ─────────────────────────────────────────────────────────────────────────────
# TAB 6 — MISSING VALUES & IMPUTATION  ← NEW
# ─────────────────────────────────────────────────────────────────────────────

with tabs[5]:

    st.markdown('<div class="section-title">🩹 Missing Values Detection & Imputation</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        missing_counts = df[num].isnull().sum()
        missing_pcts   = df[num].isnull().mean() * 100
        has_missing    = missing_counts[missing_counts > 0]

        # ── Overview Metrics
        mc1, mc2, mc3, mc4 = st.columns(4)
        mc1.markdown(f'<div class="metric-card"><h4>Total Missing Cells</h4><p>{df.isnull().sum().sum():,}</p></div>', unsafe_allow_html=True)
        mc2.markdown(f'<div class="metric-card"><h4>Columns with Missing</h4><p>{len(has_missing)}</p></div>', unsafe_allow_html=True)
        mc3.markdown(f'<div class="metric-card"><h4>Complete Rows</h4><p>{df.dropna().shape[0]:,}</p></div>', unsafe_allow_html=True)
        mc4.markdown(f'<div class="metric-card"><h4>Missing Rate</h4><p>{(df.isnull().sum().sum() / df.size * 100):.2f}%</p></div>', unsafe_allow_html=True)

        st.markdown("")

        if len(has_missing) == 0:
            st.markdown("""
            <div class="insight-box">
                ✅ <b>No missing values detected</b> in this dataset.
                All numeric columns are complete — no imputation needed before modelling.
            </div>
            """, unsafe_allow_html=True)
        else:
            col_miss_left, col_miss_right = st.columns([1.2, 1.8])

            with col_miss_left:
                st.markdown('<div class="section-title">Missing Values by Column</div>', unsafe_allow_html=True)
                miss_df = pd.DataFrame({
                    "Column"     : has_missing.index,
                    "Missing #"  : has_missing.values,
                    "Missing %"  : missing_pcts[has_missing.index].round(2).values,
                }).sort_values("Missing %", ascending=False)

                def color_miss(val):
                    if val < 5:  return "background:#e8f5e9;color:#2e7d32;"
                    elif val < 20: return "background:#fff8e1;color:#e65100;"
                    return "background:#ffebee;color:#c62828;font-weight:bold;"

                styled_miss = miss_df.style.applymap(color_miss, subset=["Missing %"])
                st.dataframe(styled_miss, use_container_width=True, height=300)

            with col_miss_right:
                st.markdown('<div class="section-title">Missing Values Heatmap</div>', unsafe_allow_html=True)
                fig, ax = plt.subplots(figsize=(8, 4))
                miss_sample = df[num].isnull().astype(int)
                if len(miss_sample) > 500:
                    miss_sample = miss_sample.sample(500, random_state=42)
                if len(has_missing) > 0:
                    sns.heatmap(
                        miss_sample[has_missing.index.tolist()],
                        ax=ax, cmap="RdYlGn_r",
                        cbar=True, yticklabels=False,
                        linewidths=0, xticklabels=True
                    )
                    ax.set_title("Missing Values Pattern (yellow = missing)", fontsize=10, weight="bold")
                    ax.set_xlabel("Columns"); ax.set_ylabel("Rows (sample)")
                    plt.xticks(rotation=35, ha="right", fontsize=8)
                else:
                    ax.text(0.5, 0.5, "No missing values", ha="center", va="center", fontsize=12)
                    ax.axis("off")
                plt.tight_layout()
                st.pyplot(fig, use_container_width=True)
                plt.close()

        st.markdown("---")

        # ── Imputation Controls
        st.markdown('<div class="section-title">Imputation Strategy</div>', unsafe_allow_html=True)

        ic1, ic2, ic3 = st.columns([1.5, 1.5, 3])

        with ic1:
            imp_col = st.selectbox(
                "Column to Impute",
                num,
                key="imp_col"
            )
        with ic2:
            imp_method = st.selectbox(
                "Imputation Method",
                ["Mean", "Median", "Mode", "KNN (k=5)", "Forward Fill", "Backward Fill", "Constant (0)"],
                key="imp_method"
            )
        with ic3:
            st.markdown("")
            st.markdown("""
            <div style="background:#f8f9ff;border-radius:8px;padding:10px 14px;
                        font-size:0.82rem;color:#333;line-height:1.7;">
                <b>Guide:</b> Mean/Median — for normally/skewed distributions<br>
                Mode — for categorical or integer columns<br>
                KNN — best quality, uses neighbouring rows<br>
                Forward/Backward Fill — for time-series data
            </div>
            """, unsafe_allow_html=True)

        col_imp_btn, col_imp_all, _ = st.columns([2, 2, 6])

        with col_imp_btn:
            imp_single = st.button(f"🩹 Impute: {imp_col}", key="btn_imp_single")

        with col_imp_all:
            imp_all_btn = st.button("🩹 Impute All (Median)", key="btn_imp_all")

        if imp_single or imp_all_btn:
            df_imp = st.session_state.df_work.copy()

            def do_impute(df_i, col, method):
                if method == "Mean":
                    df_i[col].fillna(df_i[col].mean(), inplace=True)
                elif method == "Median":
                    df_i[col].fillna(df_i[col].median(), inplace=True)
                elif method == "Mode":
                    mode_v = df_i[col].mode()
                    if len(mode_v) > 0:
                        df_i[col].fillna(mode_v[0], inplace=True)
                elif method == "KNN (k=5)":
                    imputer   = KNNImputer(n_neighbors=5)
                    df_i[col] = imputer.fit_transform(df_i[[col]])
                elif method == "Forward Fill":
                    df_i[col].fillna(method="ffill", inplace=True)
                elif method == "Backward Fill":
                    df_i[col].fillna(method="bfill", inplace=True)
                elif method == "Constant (0)":
                    df_i[col].fillna(0, inplace=True)
                return df_i

            if imp_single:
                df_imp = do_impute(df_imp, imp_col, imp_method)
                st.success(f"✅ Imputed **{imp_col}** using **{imp_method}**.")

            if imp_all_btn:
                for col in num:
                    df_imp = do_impute(df_imp, col, "Median")
                st.success(f"✅ All {len(num)} numeric columns imputed using Median.")

            st.session_state.df_imputed = df_imp.copy()
            st.session_state.df_work    = df_imp.copy()

            # Show before/after distribution for imp_col
            if imp_single and imp_col:
                fig, (a1, a2) = plt.subplots(1, 2, figsize=(10, 4))
                orig = st.session_state.df_raw[imp_col].dropna()
                after = df_imp[imp_col].dropna()

                a1.hist(orig,  bins=35, color=BLUE, alpha=0.75, edgecolor="white")
                a1.set_title(f"Before Imputation\n{imp_col}", weight="bold")
                a1.set_xlabel(imp_col); a1.set_ylabel("Frequency")

                a2.hist(after, bins=35, color=GREEN, alpha=0.75, edgecolor="white")
                a2.set_title(f"After Imputation ({imp_method})\n{imp_col}", weight="bold")
                a2.set_xlabel(imp_col); a2.set_ylabel("Frequency")

                plt.tight_layout()
                st.pyplot(fig, use_container_width=True)
                plt.close()

        # Export imputed dataset
        if st.session_state.df_imputed is not None:
            st.markdown("---")
            csv_imp = st.session_state.df_imputed.to_csv(index=False)
            st.download_button(
                "📥 Download Imputed Dataset (CSV)",
                csv_imp,
                file_name="data_imputed.csv",
                mime="text/csv",
                key="dl_imputed"
            )


# ─────────────────────────────────────────────────────────────────────────────
# TAB 7 — MULTICOLLINEARITY (VIF)  ← NEW
# ─────────────────────────────────────────────────────────────────────────────

with tabs[6]:

    st.markdown('<div class="section-title">🔗 Multicollinearity Analysis — Variance Inflation Factor (VIF)</div>', unsafe_allow_html=True)

    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        tgt = st.session_state.target_col
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        # ── Guide card
        st.markdown("""
        <div style="background:#e8eaf6;border-radius:8px;padding:14px 18px;
                    font-size:0.85rem;color:#1a237e;margin-bottom:16px;line-height:1.7;">
            <b>What is VIF?</b> The Variance Inflation Factor measures how much
            a feature's variance is inflated due to correlation with other features.<br>
            &bull; <b>VIF = 1</b> — No multicollinearity &nbsp;|&nbsp;
            <b>VIF 1–5</b> — Moderate &nbsp;|&nbsp;
            <b>VIF 5–10</b> — High &nbsp;|&nbsp;
            <b>VIF &gt; 10</b> — <span style="color:#c62828;font-weight:700;">Severe — remove or combine feature</span><br>
            Note: VIF requires at least 2 columns and no missing values.
        </div>
        """, unsafe_allow_html=True)

        # Controls
        vif_col1, vif_col2, vif_col3 = st.columns([2, 1.5, 1.5])

        with vif_col1:
            # Feature selection for VIF
            feat_options = [c for c in num if c != tgt]
            sel_features = st.multiselect(
                "Select Features for VIF Analysis",
                feat_options,
                default=st.session_state.important_vars[:10] if st.session_state.important_vars else feat_options[:8],
                key="vif_features"
            )

        with vif_col2:
            vif_threshold = st.slider("Flag VIF Above", 1.0, 20.0, 10.0, 0.5, key="vif_thresh")

        with vif_col3:
            st.markdown("")
            run_vif = st.button("🔗 Compute VIF", key="btn_run_vif", type="primary")

        if run_vif:
            if len(sel_features) < 2:
                st.error("Please select at least 2 features.")
            else:
                df_vif = df[sel_features].dropna()

                if len(df_vif) < 10:
                    st.error("Not enough complete rows for VIF computation.")
                else:
                    try:
                        # Standardise before VIF
                        scaler     = StandardScaler()
                        X_scaled   = scaler.fit_transform(df_vif)
                        X_df       = pd.DataFrame(X_scaled, columns=sel_features)

                        vif_values = [
                            variance_inflation_factor(X_df.values, i)
                            for i in range(X_df.shape[1])
                        ]

                        vif_df = pd.DataFrame({
                            "Feature"    : sel_features,
                            "VIF"        : [round(v, 3) for v in vif_values],
                            "Status"     : [
                                "🔴 Severe" if v > 10
                                else ("🟡 High" if v > 5
                                else ("🟢 Moderate" if v > 1
                                else "✅ Clean"))
                                for v in vif_values
                            ]
                        }).sort_values("VIF", ascending=False)

                        vif_df["Flagged"] = vif_df["VIF"] > vif_threshold

                        st.session_state["vif_result"] = vif_df

                    except Exception as e:
                        st.error(f"VIF computation error: {e}")

        if "vif_result" in st.session_state:
            vif_df = st.session_state["vif_result"]
            flagged = vif_df[vif_df["Flagged"]]
            safe    = vif_df[~vif_df["Flagged"]]

            # Metrics
            vc1, vc2, vc3, vc4 = st.columns(4)
            vc1.markdown(f'<div class="metric-card"><h4>Features Analysed</h4><p>{len(vif_df)}</p></div>', unsafe_allow_html=True)
            vc2.markdown(f'<div class="metric-card"><h4>High VIF (flagged)</h4><p>{len(flagged)}</p></div>', unsafe_allow_html=True)
            vc3.markdown(f'<div class="metric-card"><h4>Safe Features</h4><p>{len(safe)}</p></div>', unsafe_allow_html=True)
            vc4.markdown(f'<div class="metric-card"><h4>Max VIF</h4><p>{vif_df["VIF"].max():.2f}</p></div>', unsafe_allow_html=True)

            st.markdown("")
            col_vif_tbl, col_vif_bar = st.columns([1.2, 1.8])

            with col_vif_tbl:
                st.markdown('<div class="section-title">VIF Table</div>', unsafe_allow_html=True)

                def color_vif(val):
                    try:
                        v = float(val)
                        if v > 10:  return "background:#ffebee;color:#c62828;font-weight:bold;"
                        elif v > 5:  return "background:#fff8e1;color:#e65100;font-weight:bold;"
                        elif v > 1:  return "background:#fff9c4;color:#555;"
                        return "background:#e8f5e9;color:#2e7d32;"
                    except: return ""

                styled_vif = vif_df[["Feature", "VIF", "Status"]].style.applymap(
                    color_vif, subset=["VIF"]
                )
                st.dataframe(styled_vif, use_container_width=True, height=380)

            with col_vif_bar:
                st.markdown('<div class="section-title">VIF Bar Chart</div>', unsafe_allow_html=True)
                fig, ax = plt.subplots(figsize=(8, max(4, len(vif_df) * 0.35 + 1)))
                bar_colors = [
                    RED if v > 10 else (ORANGE if v > 5 else (BLUE if v > 1 else GREEN))
                    for v in vif_df["VIF"]
                ]
                vif_sorted = vif_df.sort_values("VIF", ascending=True)
                bars = ax.barh(vif_sorted["Feature"], vif_sorted["VIF"], color=[
                    RED if v > 10 else (ORANGE if v > 5 else (BLUE if v > 1 else GREEN))
                    for v in vif_sorted["VIF"]
                ], edgecolor="white", linewidth=0.5)
                ax.axvline(5,  color=ORANGE, linestyle="--", linewidth=1.5, label="VIF = 5")
                ax.axvline(10, color=RED,    linestyle="--", linewidth=1.8, label=f"Threshold = {vif_threshold:.0f}")
                ax.axvline(vif_threshold, color="black", linestyle=":", linewidth=1.5, alpha=0.6)

                # Add value labels
                for bar, val in zip(bars, vif_sorted["VIF"]):
                    ax.text(
                        bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                        f"{val:.2f}", va="center", ha="left", fontsize=8, color="#333"
                    )

                ax.set_title("Variance Inflation Factor by Feature", weight="bold", fontsize=11)
                ax.set_xlabel("VIF Value")
                ax.legend(fontsize=8)
                plt.tight_layout()
                st.pyplot(fig, use_container_width=True)
                plt.close()

            # Correlation heatmap between flagged features
            if len(flagged) >= 2:
                st.markdown("---")
                st.markdown('<div class="section-title">Correlation Heatmap — Flagged Features</div>', unsafe_allow_html=True)
                flag_cols = flagged["Feature"].tolist()
                fig2, ax2 = plt.subplots(figsize=(max(6, len(flag_cols) * 0.9), max(4, len(flag_cols) * 0.8)))
                corr_flag = df[flag_cols].corr()
                sns.heatmap(
                    corr_flag, ax=ax2, annot=True, fmt=".2f",
                    cmap="RdYlGn", vmin=-1, vmax=1,
                    linewidths=0.5, linecolor="#e0e0e0",
                    annot_kws={"size": 9}
                )
                ax2.set_title("Pairwise Correlations — High-VIF Features", weight="bold")
                plt.xticks(rotation=35, ha="right"); plt.yticks(rotation=0)
                plt.tight_layout()
                st.pyplot(fig2, use_container_width=True)
                plt.close()

            # Recommendation
            st.markdown("---")
            if len(flagged) > 0:
                flag_list = ", ".join(flagged["Feature"].tolist())
                st.markdown(f"""
                <div class="warning-box">
                    ⚠️ <b>Multicollinearity Detected</b><br>
                    The following features have VIF &gt; {vif_threshold:.0f}: <b>{flag_list}</b><br><br>
                    <b>Recommendations:</b><br>
                    • Consider removing the feature with the highest VIF first, then re-run.<br>
                    • Or combine highly correlated features using PCA / domain knowledge.<br>
                    • Keep only one feature from each highly correlated pair.
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div class="insight-box">
                    ✅ <b>No significant multicollinearity detected.</b>
                    All selected features have VIF below the threshold —
                    suitable for Regression modelling without dimension reduction.
                </div>
                """, unsafe_allow_html=True)

            # Export
            csv_vif = vif_df.to_csv(index=False)
            st.download_button(
                "📥 Export VIF Table (CSV)",
                csv_vif,
                file_name="vif_analysis.csv",
                mime="text/csv",
                key="dl_vif"
            )


# =============================================================================

# =============================================================================
# TAB 9 — HR KPI DASHBOARD
# =============================================================================
with tabs[8]:
    st.markdown('<div class="section-title">📦 HR KPI Dashboard</div>', unsafe_allow_html=True)
    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        # ── KPI Cards
        total_emp    = len(df)
        attr_rate    = df["Attrition_flag"].mean()*100   if "Attrition_flag" in df.columns else 0
        avg_income   = df["MonthlyIncome"].mean()         if "MonthlyIncome"  in df.columns else 0
        avg_age      = df["Age"].mean()                   if "Age"            in df.columns else 0
        avg_tenure   = df["YearsAtCompany"].mean()        if "YearsAtCompany" in df.columns else 0
        overtime_pct = df["OverTime_flag"].mean()*100     if "OverTime_flag"  in df.columns else 0

        k1,k2,k3,k4,k5,k6 = st.columns(6)
        for col_w, label, val, color in zip(
            [k1,k2,k3,k4,k5,k6],
            ["Total Employees","Attrition Rate","Avg Monthly Income",
             "Avg Age","Avg Tenure (yrs)","OverTime %"],
            [f"{total_emp:,}", f"{attr_rate:.1f}%", f"${avg_income:,.0f}",
             f"{avg_age:.1f} yrs", f"{avg_tenure:.1f} yrs", f"{overtime_pct:.1f}%"],
            ["#1565c0","#c62828","#2e7d32","#e65100","#6a1b9a","#ad1457"]
        ):
            col_w.markdown(f"""
            <div style='background:white;border-radius:10px;padding:1rem;
                        border-left:5px solid {color};box-shadow:0 2px 6px rgba(0,0,0,.08);
                        text-align:center;'>
                <p style='color:#546e7a;font-size:0.75rem;margin:0;'>{label}</p>
                <h3 style='color:{color};margin:0.3rem 0 0 0;font-size:1.1rem;'>{val}</h3>
            </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        col_l, col_r = st.columns([3,2])

        with col_l:
            st.markdown("#### 🏢 Attrition Rate by Department")
            if "Department" in df.columns and "Attrition_flag" in df.columns:
                dept = df.groupby("Department")["Attrition_flag"].agg(["mean","count"]).reset_index()
                dept.columns = ["Department","Attrition Rate","Count"]
                dept["Attrition Rate"] *= 100
                dept = dept.sort_values("Attrition Rate", ascending=False)
                fig, ax = plt.subplots(figsize=(8,3.5))
                colors_d = ["#c62828" if v>20 else "#e65100" if v>15 else "#2e7d32"
                            for v in dept["Attrition Rate"]]
                bars = ax.bar(dept["Department"], dept["Attrition Rate"],
                              color=colors_d, alpha=0.85)
                for bar, val in zip(bars, dept["Attrition Rate"]):
                    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
                            f"{val:.1f}%", ha="center", fontsize=9, fontweight="bold")
                ax.axhline(attr_rate, color="#1565c0", linestyle="--",
                           linewidth=1.5, label=f"Overall: {attr_rate:.1f}%")
                ax.set_ylabel("Attrition Rate (%)", fontsize=9)
                ax.set_title("Attrition Rate by Department", fontsize=11, fontweight="bold")
                ax.legend(fontsize=8)
                ax.grid(axis="y", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        with col_r:
            st.markdown("#### 💼 Headcount by Job Role")
            if "JobRole" in df.columns:
                role_counts = df["JobRole"].value_counts()
                fig, ax = plt.subplots(figsize=(5,3.5))
                ax.barh(role_counts.index[::-1], role_counts.values[::-1],
                        color="#1565c0", alpha=0.8)
                ax.set_xlabel("Count", fontsize=9)
                ax.set_title("Employees by Job Role", fontsize=11, fontweight="bold")
                ax.grid(axis="x", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        st.markdown("<br>", unsafe_allow_html=True)
        col_l2, col_r2 = st.columns(2)

        with col_l2:
            st.markdown("#### 📊 Age Distribution by Attrition")
            if "Age" in df.columns and "Attrition_flag" in df.columns:
                fig, ax = plt.subplots(figsize=(6,3.5))
                for val, label, color in [(1,"Left","#c62828"),(0,"Stayed","#1565c0")]:
                    data = df[df["Attrition_flag"]==val]["Age"]
                    ax.hist(data, bins=20, alpha=0.6, color=color,
                            label=f"{label} (n={len(data)})", edgecolor="white")
                ax.set_xlabel("Age", fontsize=9)
                ax.set_ylabel("Count", fontsize=9)
                ax.set_title("Age Distribution — Left vs Stayed", fontsize=11, fontweight="bold")
                ax.legend(fontsize=8)
                ax.grid(alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        with col_r2:
            st.markdown("#### ⏱️ Attrition by Years at Company")
            if "YearsAtCompany" in df.columns and "Attrition_flag" in df.columns:
                bins = [0,1,3,5,10,15,20,40]
                labels = ["<1","1-3","3-5","5-10","10-15","15-20","20+"]
                df["tenure_group"] = pd.cut(df["YearsAtCompany"], bins=bins, labels=labels)
                tg = df.groupby("tenure_group", observed=True)["Attrition_flag"].mean()*100
                fig, ax = plt.subplots(figsize=(6,3.5))
                ax.bar(tg.index, tg.values, color="#e65100", alpha=0.85)
                ax.axhline(attr_rate, color="#1565c0", linestyle="--",
                           linewidth=1.5, label=f"Overall: {attr_rate:.1f}%")
                ax.set_xlabel("Years at Company", fontsize=9)
                ax.set_ylabel("Attrition Rate (%)", fontsize=9)
                ax.set_title("Attrition Rate by Tenure", fontsize=11, fontweight="bold")
                ax.legend(fontsize=8)
                ax.grid(axis="y", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)
                df.drop(columns=["tenure_group"], inplace=True, errors="ignore")

        # ── KPI Summary Table
        st.markdown("<br>")
        st.markdown("#### 📋 HR KPI Summary")
        kpi_data = {
            "KPI": ["Total Employees","Attrition Count","Attrition Rate",
                    "Avg Monthly Income","Avg Age","Avg Tenure",
                    "OverTime %","Departments","Job Roles"],
            "Value": [
                f"{total_emp:,}",
                f"{int(df['Attrition_flag'].sum()) if 'Attrition_flag' in df.columns else 'N/A'}",
                f"{attr_rate:.1f}%",
                f"${avg_income:,.0f}",
                f"{avg_age:.1f} years",
                f"{avg_tenure:.1f} years",
                f"{overtime_pct:.1f}%",
                f"{df['Department'].nunique() if 'Department' in df.columns else 'N/A'}",
                f"{df['JobRole'].nunique() if 'JobRole' in df.columns else 'N/A'}",
            ]
        }
        st.dataframe(pd.DataFrame(kpi_data), use_container_width=True, hide_index=True)
        st.download_button("📥 Export KPI Table (CSV)",
                           pd.DataFrame(kpi_data).to_csv(index=False),
                           "hr_kpi.csv", "text/csv", key="dl_hr_kpi")


# =============================================================================
# TAB 10 — ATTRITION DEEP DIVE
# =============================================================================
with tabs[9]:
    st.markdown('<div class="section-title">🔍 Attrition Deep Dive</div>', unsafe_allow_html=True)
    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        if "Attrition_flag" not in df.columns:
            st.warning("Attrition_flag column not found.")
        else:
            attr_rate = df["Attrition_flag"].mean()*100

            st.markdown("""
            <div style='background:#fce4ec;border-radius:10px;padding:1rem 1.2rem;margin-bottom:1rem;'>
            Attrition Deep Dive identifies <b>who leaves, why, and what factors predict departure</b>.
            Every chart below is a direct business insight for HR retention strategy.
            </div>""", unsafe_allow_html=True)

            # ── Row 1: Satisfaction scores vs attrition
            st.markdown("#### 😊 Satisfaction Scores — Left vs Stayed")
            sat_cols = [c for c in ["JobSatisfaction","EnvironmentSatisfaction",
                                     "RelationshipSatisfaction","WorkLifeBalance"]
                        if c in df.columns]
            if sat_cols:
                fig, axes = plt.subplots(1, len(sat_cols), figsize=(14, 4))
                if len(sat_cols)==1: axes=[axes]
                for ax, col in zip(axes, sat_cols):
                    for val, label, color in [(1,"Left","#c62828"),(0,"Stayed","#1565c0")]:
                        grp = df[df["Attrition_flag"]==val][col].value_counts().sort_index()
                        ax.bar(grp.index + (0.2 if val==1 else -0.2),
                               grp.values, width=0.35,
                               color=color, alpha=0.8, label=label)
                    ax.set_title(col.replace("Satisfaction","").replace("Balance","Bal"),
                                 fontsize=9, fontweight="bold")
                    ax.set_xlabel("Score (1-4)", fontsize=8)
                    ax.legend(fontsize=7)
                    ax.grid(axis="y", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

            st.markdown("<br>")

            # ── Row 2: BusinessTravel + OverTime vs attrition
            col_l, col_r = st.columns(2)
            with col_l:
                st.markdown("#### ✈️ Attrition by Business Travel")
                if "BusinessTravel" in df.columns:
                    bt = df.groupby("BusinessTravel")["Attrition_flag"].mean()*100
                    bt = bt.sort_values(ascending=False)
                    fig, ax = plt.subplots(figsize=(6,3.5))
                    colors_bt = ["#c62828","#e65100","#2e7d32"]
                    ax.bar(bt.index, bt.values, color=colors_bt[:len(bt)], alpha=0.85)
                    for i,(idx,val) in enumerate(bt.items()):
                        ax.text(i, val+0.3, f"{val:.1f}%", ha="center",
                                fontsize=9, fontweight="bold")
                    ax.axhline(attr_rate, color="#1565c0", linestyle="--",
                               linewidth=1.5, label=f"Overall: {attr_rate:.1f}%")
                    ax.set_ylabel("Attrition Rate (%)", fontsize=9)
                    ax.set_title("Attrition by Business Travel", fontsize=11, fontweight="bold")
                    ax.legend(fontsize=8)
                    ax.grid(axis="y", alpha=0.3)
                    fig.tight_layout()
                    st.pyplot(fig); plt.close(fig)

            with col_r:
                st.markdown("#### ⏰ Attrition by OverTime")
                if "OverTime_flag" in df.columns:
                    ot = df.groupby("OverTime_flag")["Attrition_flag"].mean()*100
                    fig, ax = plt.subplots(figsize=(5,3.5))
                    ax.bar(["No OverTime","OverTime"], ot.values,
                           color=["#2e7d32","#c62828"], alpha=0.85, width=0.5)
                    for i,val in enumerate(ot.values):
                        ax.text(i, val+0.3, f"{val:.1f}%", ha="center",
                                fontsize=11, fontweight="bold")
                    ax.set_ylabel("Attrition Rate (%)", fontsize=9)
                    ax.set_title("Attrition: OverTime vs No OverTime",
                                 fontsize=11, fontweight="bold")
                    ax.grid(axis="y", alpha=0.3)
                    fig.tight_layout()
                    st.pyplot(fig); plt.close(fig)

            st.markdown("<br>")

            # ── Row 3: JobRole attrition + JobLevel attrition
            col_l2, col_r2 = st.columns([3,2])
            with col_l2:
                st.markdown("#### 💼 Attrition Rate by Job Role")
                if "JobRole" in df.columns:
                    role_attr = (df.groupby("JobRole")["Attrition_flag"]
                                   .mean()*100).sort_values(ascending=False)
                    fig, ax = plt.subplots(figsize=(9,4))
                    colors_r = ["#c62828" if v>20 else "#e65100" if v>15
                                else "#2e7d32" for v in role_attr.values]
                    ax.barh(role_attr.index[::-1], role_attr.values[::-1],
                            color=colors_r[::-1], alpha=0.85)
                    ax.axvline(attr_rate, color="#1565c0", linestyle="--",
                               linewidth=1.5, label=f"Overall: {attr_rate:.1f}%")
                    ax.set_xlabel("Attrition Rate (%)", fontsize=9)
                    ax.set_title("Attrition Rate by Job Role", fontsize=11, fontweight="bold")
                    ax.legend(fontsize=8)
                    ax.grid(axis="x", alpha=0.3)
                    fig.tight_layout()
                    st.pyplot(fig); plt.close(fig)

            with col_r2:
                st.markdown("#### 📶 Attrition by Job Level")
                if "JobLevel" in df.columns:
                    jl = df.groupby("JobLevel")["Attrition_flag"].mean()*100
                    fig, ax = plt.subplots(figsize=(5,4))
                    ax.bar(jl.index, jl.values, color="#6a1b9a", alpha=0.85)
                    ax.axhline(attr_rate, color="#1565c0", linestyle="--",
                               linewidth=1.5, label=f"Overall: {attr_rate:.1f}%")
                    ax.set_xlabel("Job Level (1=Junior → 5=Senior)", fontsize=8)
                    ax.set_ylabel("Attrition Rate (%)", fontsize=9)
                    ax.set_title("Attrition by Job Level", fontsize=11, fontweight="bold")
                    ax.legend(fontsize=8)
                    ax.grid(axis="y", alpha=0.3)
                    fig.tight_layout()
                    st.pyplot(fig); plt.close(fig)

            # ── Attrition Summary Table
            st.markdown("<br>")
            st.markdown("#### 📋 Attrition Summary by Department & Role")
            if "Department" in df.columns and "JobRole" in df.columns:
                summary = df.groupby(["Department","JobRole"]).agg(
                    Total        = ("Attrition_flag","count"),
                    Left         = ("Attrition_flag","sum"),
                    Attrition_Pct= ("Attrition_flag","mean"),
                    Avg_Income   = ("MonthlyIncome","mean"),
                    Avg_Age      = ("Age","mean"),
                ).round(2).reset_index()
                summary["Attrition_Pct"] = (summary["Attrition_Pct"]*100).round(1)
                summary["Avg_Income"]    = summary["Avg_Income"].round(0)
                summary["Avg_Age"]       = summary["Avg_Age"].round(1)
                summary.columns = ["Department","Job Role","Total","Left",
                                   "Attrition %","Avg Income ($)","Avg Age"]
                summary = summary.sort_values("Attrition %", ascending=False)
                st.dataframe(summary, use_container_width=True, hide_index=True)
                st.download_button("📥 Export Attrition Table",
                                   summary.to_csv(index=False),
                                   "attrition_summary.csv","text/csv",
                                   key="dl_attr_summary")


# =============================================================================
# TAB 11 — SALARY ANALYSIS
# =============================================================================
with tabs[10]:
    st.markdown('<div class="section-title">💰 Salary Analysis</div>', unsafe_allow_html=True)
    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        # ── Salary KPI cards
        avg_sal  = df["MonthlyIncome"].mean()   if "MonthlyIncome" in df.columns else 0
        med_sal  = df["MonthlyIncome"].median() if "MonthlyIncome" in df.columns else 0
        min_sal  = df["MonthlyIncome"].min()    if "MonthlyIncome" in df.columns else 0
        max_sal  = df["MonthlyIncome"].max()    if "MonthlyIncome" in df.columns else 0

        s1,s2,s3,s4 = st.columns(4)
        for col_w, label, val, color in zip(
            [s1,s2,s3,s4],
            ["Avg Monthly Income","Median Income","Min Income","Max Income"],
            [f"${avg_sal:,.0f}", f"${med_sal:,.0f}",
             f"${min_sal:,.0f}", f"${max_sal:,.0f}"],
            ["#1565c0","#2e7d32","#e65100","#6a1b9a"]
        ):
            col_w.markdown(f"""
            <div style='background:white;border-radius:10px;padding:1rem;
                        border-left:5px solid {color};text-align:center;'>
                <p style='color:#546e7a;font-size:0.75rem;margin:0;'>{label}</p>
                <h3 style='color:{color};margin:0.3rem 0 0 0;'>{val}</h3>
            </div>""", unsafe_allow_html=True)

        st.markdown("<br>")

        col_l, col_r = st.columns(2)

        with col_l:
            st.markdown("#### 💵 Income Distribution by Attrition")
            if "MonthlyIncome" in df.columns and "Attrition_flag" in df.columns:
                fig, ax = plt.subplots(figsize=(6,4))
                for val, label, color in [(0,"Stayed","#1565c0"),(1,"Left","#c62828")]:
                    data = df[df["Attrition_flag"]==val]["MonthlyIncome"]
                    ax.hist(data, bins=30, alpha=0.6, color=color,
                            label=f"{label} (avg: ${data.mean():,.0f})",
                            edgecolor="white")
                ax.set_xlabel("Monthly Income ($)", fontsize=9)
                ax.set_ylabel("Count", fontsize=9)
                ax.set_title("Income Distribution — Left vs Stayed",
                             fontsize=11, fontweight="bold")
                ax.legend(fontsize=8)
                ax.grid(alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        with col_r:
            st.markdown("#### 📶 Avg Income by Job Level")
            if "JobLevel" in df.columns and "MonthlyIncome" in df.columns:
                jl_income = df.groupby("JobLevel")["MonthlyIncome"].mean()
                fig, ax = plt.subplots(figsize=(6,4))
                bars = ax.bar(jl_income.index, jl_income.values,
                              color="#1565c0", alpha=0.85)
                for bar, val in zip(bars, jl_income.values):
                    ax.text(bar.get_x()+bar.get_width()/2,
                            bar.get_height()+50,
                            f"${val:,.0f}", ha="center", fontsize=8)
                ax.set_xlabel("Job Level", fontsize=9)
                ax.set_ylabel("Avg Monthly Income ($)", fontsize=9)
                ax.set_title("Avg Income by Job Level", fontsize=11, fontweight="bold")
                ax.grid(axis="y", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        st.markdown("<br>")
        col_l2, col_r2 = st.columns(2)

        with col_l2:
            st.markdown("#### 🏢 Avg Income by Department")
            if "Department" in df.columns and "MonthlyIncome" in df.columns:
                dept_inc = df.groupby("Department")["MonthlyIncome"].mean().sort_values()
                fig, ax = plt.subplots(figsize=(6,3.5))
                ax.barh(dept_inc.index, dept_inc.values,
                        color="#2e7d32", alpha=0.85)
                for i, val in enumerate(dept_inc.values):
                    ax.text(val+50, i, f"${val:,.0f}", va="center", fontsize=9)
                ax.set_xlabel("Avg Monthly Income ($)", fontsize=9)
                ax.set_title("Avg Income by Department", fontsize=11, fontweight="bold")
                ax.grid(axis="x", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        with col_r2:
            st.markdown("#### 📈 Salary Hike vs Performance Rating")
            if "PerformanceRating" in df.columns and "PercentSalaryHike" in df.columns:
                pr_hike = df.groupby("PerformanceRating")["PercentSalaryHike"].mean()
                fig, ax = plt.subplots(figsize=(6,3.5))
                ax.bar(pr_hike.index, pr_hike.values,
                       color="#e65100", alpha=0.85)
                for i, val in enumerate(pr_hike.values):
                    ax.text(pr_hike.index[i], val+0.1,
                            f"{val:.1f}%", ha="center", fontsize=9, fontweight="bold")
                ax.set_xlabel("Performance Rating", fontsize=9)
                ax.set_ylabel("Avg Salary Hike (%)", fontsize=9)
                ax.set_title("Salary Hike by Performance Rating",
                             fontsize=11, fontweight="bold")
                ax.grid(axis="y", alpha=0.3)
                fig.tight_layout()
                st.pyplot(fig); plt.close(fig)

        # ── Salary by Role table
        st.markdown("<br>")
        st.markdown("#### 📋 Salary Summary by Job Role")
        if "JobRole" in df.columns and "MonthlyIncome" in df.columns:
            sal_tbl = df.groupby("JobRole")["MonthlyIncome"].agg(
                ["mean","median","min","max","count"]).round(0).reset_index()
            sal_tbl.columns = ["Job Role","Avg ($)","Median ($)",
                                "Min ($)","Max ($)","Count"]
            sal_tbl = sal_tbl.sort_values("Avg ($)", ascending=False)
            st.dataframe(sal_tbl, use_container_width=True, hide_index=True)
            st.download_button("📥 Export Salary Table",
                               sal_tbl.to_csv(index=False),
                               "salary_analysis.csv","text/csv",
                               key="dl_salary")


# =============================================================================
# TAB 12 — STATISTICAL TESTS
# =============================================================================
with tabs[11]:
    st.markdown('<div class="section-title">🧪 Statistical Tests</div>', unsafe_allow_html=True)
    if st.session_state.df_raw is None:
        st.warning("Please load a dataset first.")
    else:
        df  = st.session_state.df_work
        num = df.select_dtypes(include="number").columns.tolist()
        cat = df.select_dtypes(include="object").columns.tolist()
        st.session_state.num_cols = num
        st.session_state.cat_cols = cat

        st.markdown("""
        <div style='background:#e3f2fd;border-radius:10px;padding:1rem 1.2rem;margin-bottom:1rem;'>
        Statistical tests confirm whether observed differences are <b>real or due to chance</b>.
        P-value &lt; 0.05 = statistically significant difference.
        </div>""", unsafe_allow_html=True)

        p_val = t_pval = chi_p = fp = 0.0  # defaults

        # ── TEST 1: T-Test — Income: Left vs Stayed
        st.markdown("### 📊 Test 1 — T-Test: Does Income Differ Between Left & Stayed?")
        if "MonthlyIncome" in df.columns and "Attrition_flag" in df.columns:
            g1 = df[df["Attrition_flag"]==1]["MonthlyIncome"].dropna()
            g0 = df[df["Attrition_flag"]==0]["MonthlyIncome"].dropna()
            t_stat, t_pval = stats.ttest_ind(g1, g0, equal_var=False)
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("Left Avg Income",   f"${g1.mean():,.0f}")
            c2.metric("Stayed Avg Income", f"${g0.mean():,.0f}")
            c3.metric("T-Statistic",        f"{t_stat:.2f}")
            c4.metric("P-Value",            f"{t_pval:.4f}")
            if t_pval < 0.05:
                st.success("✅ Significant difference in income — lower-paid employees leave more.")
            else:
                st.info("ℹ️ No significant income difference between groups.")

        st.markdown("---")

        # ── TEST 2: ANOVA — Does dept affect income?
        st.markdown("### 📊 Test 2 — ANOVA: Does Department Affect Monthly Income?")
        if "Department" in df.columns and "MonthlyIncome" in df.columns:
            depts  = df["Department"].unique()
            groups = [df[df["Department"]==d]["MonthlyIncome"].dropna().values
                      for d in depts]
            f_stat, p_val = stats.f_oneway(*groups)
            ca,cb,cc = st.columns(3)
            ca.metric("F-Statistic", f"{f_stat:.2f}")
            cb.metric("P-Value",     f"{p_val:.4f}")
            cc.metric("Result", "✅ Significant" if p_val<0.05 else "❌ Not Significant")
            if p_val < 0.05:
                st.success("✅ Department significantly affects income (p < 0.05).")

        st.markdown("---")

        # ── TEST 3: Chi-Square — OverTime vs Attrition
        st.markdown("### 📊 Test 3 — Chi² Test: Is OverTime Related to Attrition?")
        if "OverTime_flag" in df.columns and "Attrition_flag" in df.columns:
            ct = pd.crosstab(df["OverTime_flag"], df["Attrition_flag"])
            chi2, chi_p, dof, _ = stats.chi2_contingency(ct)
            d1,d2,d3,d4 = st.columns(4)
            d1.metric("Chi² Statistic",      f"{chi2:.2f}")
            d2.metric("P-Value",             f"{chi_p:.4f}")
            d3.metric("Degrees of Freedom",  f"{dof}")
            d4.metric("Result", "✅ Dependent" if chi_p<0.05 else "❌ Independent")
            if chi_p < 0.05:
                st.success("✅ OverTime and Attrition are strongly related (p < 0.05).")
            ct_pct = ct.div(ct.sum(axis=1), axis=0)*100
            fig, ax = plt.subplots(figsize=(6,3.5))
            ct_pct.plot(kind="bar", stacked=True, ax=ax,
                        color=["#2e7d32","#c62828"], alpha=0.85)
            ax.set_xticklabels(["No OverTime","OverTime"], rotation=0)
            ax.set_ylabel("Percentage (%)", fontsize=9)
            ax.set_title("Attrition Rate by OverTime Status",
                         fontsize=11, fontweight="bold")
            ax.legend(["Stayed (0)","Left (1)"], fontsize=9)
            ax.grid(axis="y", alpha=0.3)
            fig.tight_layout()
            st.pyplot(fig); plt.close(fig)

        st.markdown("---")

        # ── TEST 4: ANOVA — JobSatisfaction vs Attrition groups
        st.markdown("### 📊 Test 4 — ANOVA: Does Job Satisfaction Differ by Attrition?")
        if "JobSatisfaction" in df.columns and "Attrition_flag" in df.columns:
            gs1 = df[df["Attrition_flag"]==1]["JobSatisfaction"].dropna().values
            gs0 = df[df["Attrition_flag"]==0]["JobSatisfaction"].dropna().values
            fs, fp = stats.f_oneway(gs1, gs0)
            e1,e2,e3 = st.columns(3)
            e1.metric("F-Statistic", f"{fs:.2f}")
            e2.metric("P-Value",     f"{fp:.4f}")
            e3.metric("Result", "✅ Significant" if fp<0.05 else "❌ Not Significant")
            if fp < 0.05:
                st.success("✅ Job satisfaction significantly differs — dissatisfied employees leave more.")

        # ── Summary table
        st.markdown("<br>")
        st.markdown("#### 📋 Statistical Tests Summary")
        tests_df = pd.DataFrame({
            "Test"       : ["T-Test (Welch)","ANOVA","Chi-Square","ANOVA"],
            "Question"   : [
                "Does income differ: Left vs Stayed?",
                "Does department affect income?",
                "Is OverTime related to Attrition?",
                "Does job satisfaction differ by Attrition?",
            ],
            "P-Value"    : [f"{t_pval:.4f}", f"{p_val:.4f}",
                            f"{chi_p:.4f}",  f"{fp:.4f}"],
            "Significant": [
                "✅ Yes" if t_pval<0.05 else "❌ No",
                "✅ Yes" if p_val <0.05 else "❌ No",
                "✅ Yes" if chi_p <0.05 else "❌ No",
                "✅ Yes" if fp    <0.05 else "❌ No",
            ]
        })
        st.dataframe(tests_df, use_container_width=True, hide_index=True)


# H — FOOTER
# =============================================================================

st.markdown("""
<div class="footer">
    HR Attrition · IBM HR Analytics EDA Dashboard &nbsp;|&nbsp;
    Built with Streamlit · Tabs 1–12 Complete &nbsp;|&nbsp;
    M3 · Data Analysis Portfolio
</div>
""", unsafe_allow_html=True)