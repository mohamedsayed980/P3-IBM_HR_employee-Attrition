# =============================================================================
# ML_Models_Dashboard.py
# =============================================================================
# Stage 2 — ML Models Engine (Tabs 9 → 12+1)
# Part of the End-to-End ML Engine project
# Author  : Mohamed (Mechanical Engineer / Data Analyst)
# Version : 2.0  —  May 2025
#
# Tabs:
#   Tab  9  — Regression Models       (6 models + CV + feature importance)
#   Tab 10  — Classification Models   (6 models + CV + feature importance)
#   Tab 11  — Comparison & Reports    (side-by-side + PDF/Word export)
#   Tab 12  — Predict New Data        (single-row + batch upload)
#
# Shared state with ML_EDA_Dashboard.py via st.session_state
# Run standalone:  streamlit run ML_Models.py
# =============================================================================
## path = streamlit run "E:\FINAL PROJECTS\P4_Superstore (2015-2018)\ML_Models.py" 
# ─────────────────────────────────────────────────────────────────────────────
# A  IMPORTS
# ─────────────────────────────────────────────────────────────────────────────
# A1 — Core
import streamlit as st
from fpdf import FPDF
import pandas as pd
import numpy as np
import warnings, io, datetime, time, os, concurrent.futures
warnings.filterwarnings("ignore")

# A2 — Visualisation
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import seaborn as sns
from matplotlib.patches import FancyBboxPatch

# A3 — Sklearn — Regression
from sklearn.linear_model    import LinearRegression, Ridge, Lasso
from sklearn.tree            import DecisionTreeRegressor
from sklearn.ensemble        import RandomForestRegressor, GradientBoostingRegressor

# A4 — Sklearn — Classification
from sklearn.linear_model    import LogisticRegression
from sklearn.neighbors       import KNeighborsClassifier
from sklearn.tree            import DecisionTreeClassifier
from sklearn.ensemble        import RandomForestClassifier, GradientBoostingClassifier
from sklearn.svm             import SVC

# A5 — Sklearn — Utilities
from sklearn.model_selection import (train_test_split, cross_val_score,
                                     StratifiedKFold, KFold, learning_curve)
from sklearn.preprocessing   import StandardScaler, LabelEncoder
from sklearn.metrics         import (mean_absolute_error, mean_squared_error,
                                     r2_score, accuracy_score, f1_score,
                                     precision_score, recall_score,
                                     confusion_matrix, classification_report,
                                     roc_auc_score, roc_curve)
from sklearn.inspection      import permutation_importance
import joblib, pickle

# A6 — Report generation
try:
    from reportlab.lib.pagesizes   import letter, A4
    from reportlab.lib             import colors
    from reportlab.lib.styles      import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus        import (SimpleDocTemplate, Paragraph, Spacer,
                                           Table, TableStyle, HRFlowable)
    from reportlab.lib.units       import inch
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx                      import Document
    from docx.shared               import Inches, Pt, RGBColor
    from docx.enum.text            import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ─────────────────────────────────────────────────────────────────────────────
# B  PAGE CONFIG & GLOBAL STYLE------>> in  Home.py  only 
# ─────────────────────────────────────────────────────────────────────────────
# ADD LOGO TO DASHBOARD 
import pathlib
LOGO = pathlib.Path(__file__).parent.parent / "3M_logo.png" 

#==============================================================================

# ── Color palette — complete ──────────────────────────────────────────────────
CLR = {
    "primary"   : "#1565c0",   # blue
    "success"   : "#2e7d32",   # green
    "warning"   : "#e65100",   # orange
    "danger"    : "#c62828",   # red
    "teal"      : "#00695c",   # teal
    "accent"    : "#00695c",   # teal (alias)
    "secondary" : "#455a64",   # grey-blue
    "light"     : "#e3f2fd",   # light blue
    "dark"      : "#1a237e",   # dark navy
    "purple"    : "#6a1b9a",   # purple
    "amber"     : "#f57f17",   # amber
    "pink"      : "#ad1457",   # pink
    "indigo"    : "#283593",   # indigo
    "cyan"      : "#00838f",   # cyan
    "lime"      : "#558b2f",   # lime
    "brown"     : "#4e342e",   # brown
    "grey"      : "#546e7a",   # grey
    "white"     : "#ffffff",   # white
    "black"     : "#212121",   # near black
}

# ─────────────────────────────────────────────────────────────────────────────
# C  SESSION STATE INITIALISATION
# ─────────────────────────────────────────────────────────────────────────────
def init_state():
    defaults = {
        # ── EDA hand-off (written by ML_EDA_Dashboard) ──
        "df_clean"        : None,
        "df_work"         : None,   # Stage 1 writes here
        "df_raw"          : None,   # Stage 1 raw data
        "df_original"     : None,
        "target_col"      : "MonthlyIncome",
        "feat_names"      : [],
        "num_cols"        : [],
        "cat_cols"        : [],
        "important_vars"  : [],
        "file_name"       : None,
        "insights_text"   : "",
        "final_report_text": "",
        
        # ── Internal data splits ──
        "X_train_r"  : None, "X_test_r"  : None,
        "y_train_r"  : None, "y_test_r"  : None,
        "X_train_c"  : None, "X_test_c"  : None,
        "y_train_c"  : None, "y_test_c"  : None,
        "scaler_r"   : None, "scaler_c"  : None,
        "le"         : None,
        # ── Trained model stores ──
        "reg_models"   : {},   # {name: fitted_estimator}
        "cls_models"   : {},
        "reg_results"  : {},   # {name: {r2, mae, rmse, cv_r2, …}}
        "cls_results"  : {},   # {name: {acc, f1, …}}
        "best_reg_name": None,
        "best_cls_name": None,
        # ── Tab 12 ──
        "batch_results"   : None,
        "stage2_insights" : "",
        # ── Flags ──
        "data_prepared_r" : False,
        "data_prepared_c" : False,
        "price_bins"      : [0, 50, 100, 200, 10000],
        "price_labels"    : ["Low","Medium","High","Premium"],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()
S = st.session_state   # shorthand


# ─────────────────────────────────────────────────────────────────────────────
# D  HELPER UTILITIES & helper function (for parallel training )
# ─────────────────────────────────────────────────────────────────────────────
#-----------------------------------------------------
# ── CPU info helper ──────────────────────────────────────────────────────────
import os, psutil

def get_cpu_info(use_parallel: bool, n_jobs: int = 1) -> dict:
    total   = os.cpu_count()
    used    = n_jobs if use_parallel else 1
    percent = psutil.cpu_percent(interval=0.3)
    return {"total": total, "used": used, "percent": percent}

#------------------------------------------------------

# ── D1: colour map for category badges ─────────────────────────────────────
CAT_STYLE = {
    "Low"    : ("badge-green",  "🟢"),
    "Medium" : ("badge-blue",   "🔵"),
    "High"   : ("badge-amber",  "🟡"),
    "Luxury" : ("badge-red",    "🔴"),
}

def cat_badge(cat: str) -> str:
    cls, icon = CAT_STYLE.get(cat, ("badge-purple", "⬜"))
    return f'<span class="badge {cls}">{icon} {cat}</span>'

# ── D2: small metric card ───────────────────────────────────────────────────
def metric_card(label: str, value, sub: str = "", colour: str = "") -> str:
    cls = f"metric-card {colour}" if colour else "metric-card"
    return (f'<div class="{cls}">'
            f'<div class="metric-label">{label}</div>'
            f'<div class="metric-value">{value}</div>'
            f'<div class="metric-sub">{sub}</div>'
            f'</div>')

# ── D3: section header ──────────────────────────────────────────────────────
def section(title: str):
    st.markdown(f'<div class="section-header"><h3>{title}</h3></div>',
                unsafe_allow_html=True)

def load_data() -> pd.DataFrame | None:
    """Returns cleaned dataframe — from session_state (Stage 1 or upload)."""
    # Try Stage 1 key first, then Stage 2 upload key
    df = S.get("df_work")
    if df is None:
        df = S.get("df_clean")
    if df is not None and not isinstance(df, pd.DataFrame):
        return None
    return df if (df is not None and not df.empty) else None

# ── D5: prepare regression split ────────────────────────────────────────────
def prepare_regression(df: pd.DataFrame):
    target   = S["target_col"]
    feats    = S["feat_names"] or [c for c in df.columns
                                   if c != target and
                                   pd.api.types.is_numeric_dtype(df[c])]
    S["feat_names"] = feats
    # Drop rows where target is null (e.g. delivery_days nulls for non-delivered)
    df_r = df.dropna(subset=[target]).copy()
    X = df_r[feats].fillna(df_r[feats].median())
    y = df_r[target]
    X_tr, X_te, y_tr, y_te = train_test_split(X, y, test_size=.2,
                                               random_state=42)
    sc = StandardScaler()
    X_tr_s = sc.fit_transform(X_tr)
    X_te_s = sc.transform(X_te)
    S.update({"X_train_r": X_tr_s, "X_test_r": X_te_s,
               "y_train_r": y_tr,   "y_test_r": y_te,
               "scaler_r": sc,      "data_prepared_r": True})

# ── D6: prepare classification split ────────────────────────────────────────
def prepare_classification(df: pd.DataFrame):
    # ── Detect binary target: Attrition_flag (P3) or is_satisfied (P2)
    feats   = S["feat_names"]
    df2     = df.copy()

    if "Attrition_flag" in df2.columns:
        df2.dropna(subset=["Attrition_flag"], inplace=True)
        y_enc = df2["Attrition_flag"].astype(int).values
        le    = LabelEncoder()
        le.classes_ = np.array([0, 1])
    elif "is_satisfied" in df2.columns:
        df2.dropna(subset=["is_satisfied"], inplace=True)
        y_enc = df2["is_satisfied"].astype(int).values
        le    = LabelEncoder()
        le.classes_ = np.array([0, 1])
    else:
        bins   = S["price_bins"]
        labels = S["price_labels"]
        df2["price_category"] = pd.cut(df2[S["target_col"]], bins=bins,
                                       labels=labels, right=True)
        df2.dropna(subset=["price_category"], inplace=True)
        le    = LabelEncoder()
        y_enc = le.fit_transform(df2["price_category"].astype(str))

    X      = df2[feats].fillna(df2[feats].median())
    X_tr, X_te, y_tr, y_te = train_test_split(X, y_enc, test_size=.2,
                                               random_state=42, stratify=y_enc)
    sc = StandardScaler()
    X_tr_s = sc.fit_transform(X_tr)
    X_te_s = sc.transform(X_te)
    S.update({"X_train_c": X_tr_s, "X_test_c": X_te_s,
               "y_train_c": y_tr,   "y_test_c": y_te,
               "scaler_c": sc,      "le": le,
               "data_prepared_c": True})

# ── D7: regression metrics ───────────────────────────────────────────────────
def reg_metrics(model, X_tr, X_te, y_tr, y_te, n_cv=5):
    model.fit(X_tr, y_tr)
    y_pr  = model.predict(X_te)
    cv    = cross_val_score(model, X_tr, y_tr, cv=KFold(n_cv, shuffle=True,
                            random_state=42), scoring="r2")
    return {
        "r2"     : r2_score(y_te, y_pr),
        "mae"    : mean_absolute_error(y_te, y_pr),
        "rmse"   : np.sqrt(mean_squared_error(y_te, y_pr)),
        "cv_mean": cv.mean(),
        "cv_std" : cv.std(),
        "y_pred" : y_pr,
        "model"  : model,
    }

# ── D8: classification metrics ───────────────────────────────────────────────
def cls_metrics(model, X_tr, X_te, y_tr, y_te, n_cv=5):
    model.fit(X_tr, y_tr)
    y_pr  = model.predict(X_te)
    cv    = cross_val_score(model, X_tr, y_tr,
                            cv=StratifiedKFold(n_cv, shuffle=True,
                                               random_state=42),
                            scoring="accuracy")
    res = {
        "acc"      : accuracy_score(y_te, y_pr),
        "f1"       : f1_score(y_te, y_pr, average="weighted", zero_division=0),
        "precision": precision_score(y_te, y_pr, average="weighted",
                                     zero_division=0),
        "recall"   : recall_score(y_te, y_pr, average="weighted",
                                  zero_division=0),
        "cv_mean"  : cv.mean(),
        "cv_std"   : cv.std(),
        "y_pred"   : y_pr,
        "cm"       : confusion_matrix(y_te, y_pr),
        "model"    : model,
    }
    if hasattr(model, "predict_proba"):
        try:
            y_prob = model.predict_proba(X_te)
            res["y_prob"] = y_prob
        except Exception:
            pass
    return res

# ── D9: parallel training helper ─────────────────────────────────────────────
def _train_one_reg(args):
    name, model, X_tr, X_te, y_tr, y_te = args
    try:
        r = reg_metrics(model, X_tr, X_te, y_tr, y_te)
        return name, r, None
    except Exception as e:
        return name, None, str(e)

def _train_one_cls(args):
    name, model, X_tr, X_te, y_tr, y_te = args
    try:
        r = cls_metrics(model, X_tr, X_te, y_tr, y_te)
        return name, r, None
    except Exception as e:
        return name, None, str(e)

# ── D10: feature importance chart ────────────────────────────────────────────
def plot_feature_importance(model, feat_names, title="Feature Importance",
                             colour=CLR["primary"]):
    fig, ax = plt.subplots(figsize=(7, 4))
    if hasattr(model, "feature_importances_"):
        imp = model.feature_importances_
    elif hasattr(model, "coef_"):
        imp = np.abs(model.coef_.ravel())
    else:
        ax.text(.5, .5, "Not available for this model",
                ha="center", va="center", transform=ax.transAxes)
        return fig
    idx  = np.argsort(imp)[-15:]          # top 15
    bars = ax.barh([feat_names[i] for i in idx], imp[idx],
                   color=colour, edgecolor="white", height=.65)
    ax.bar_label(bars, fmt="%.3f", padding=3, fontsize=8)
    ax.set_xlabel("Importance", fontsize=10)
    ax.set_title(title, fontsize=12, fontweight="bold", pad=10)
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return fig

# ── D11: learning curve ──────────────────────────────────────────────────────
def plot_learning_curve(model, X, y, task="reg", title="Learning Curve"):
    cv      = KFold(5, shuffle=True, random_state=42) if task == "reg" \
              else StratifiedKFold(5, shuffle=True, random_state=42)
    scoring = "r2" if task == "reg" else "accuracy"
    sizes   = np.linspace(.1, 1.0, 8)

    # learning_curve returns (train_sizes_abs, train_scores, val_scores)
    # shapes: (n_train_sizes,)  and  (n_train_sizes, n_cv_splits)
    train_sizes_abs, train_s, val_s = learning_curve(
        model, X, y,
        cv=cv, scoring=scoring,
        train_sizes=sizes, n_jobs=-1
    )

    # mean / std across CV splits → axis=1  → shape (n_train_sizes,)
    train_mean = train_s.mean(axis=1)
    train_std  = train_s.std(axis=1)
    val_mean   = val_s.mean(axis=1)
    val_std    = val_s.std(axis=1)

    fig, ax = plt.subplots(figsize=(7, 4))

    ax.fill_between(train_sizes_abs,
                    train_mean - train_std,
                    train_mean + train_std,
                    alpha=.15, color=CLR["primary"])
    ax.fill_between(train_sizes_abs,
                    val_mean - val_std,
                    val_mean + val_std,
                    alpha=.15, color=CLR["success"])
    ax.plot(train_sizes_abs, train_mean, "o-",
            color=CLR["primary"],  label="Training",   lw=2)
    ax.plot(train_sizes_abs, val_mean,   "s--",
            color=CLR["success"], label="Validation",  lw=2)

    ax.set_xlabel("Training samples", fontsize=10)
    ax.set_ylabel(scoring.upper(), fontsize=10)
    ax.set_title(title, fontsize=12, fontweight="bold", pad=10)
    ax.legend(fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    return fig

# ── D12: confusion matrix heatmap ────────────────────────────────────────────
def plot_cm(cm, labels, title="Confusion Matrix"):
    fig, ax = plt.subplots(figsize=(5, 4))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues",
                xticklabels=labels, yticklabels=labels,
                linewidths=.5, ax=ax, cbar=False)
    ax.set_xlabel("Predicted", fontsize=10)
    ax.set_ylabel("Actual",    fontsize=10)
    ax.set_title(title, fontsize=12, fontweight="bold", pad=10)
    fig.tight_layout()
    return fig

# ── D13: residual plot ───────────────────────────────────────────────────────
def plot_residuals(y_true, y_pred, title="Residuals"):
    res  = y_true - y_pred
    fig, axes = plt.subplots(1, 2, figsize=(10, 4))
    axes[0].scatter(y_pred, res, alpha=.4, color=CLR["primary"],
                    edgecolors="white", s=18)
    axes[0].axhline(0, color=CLR["danger"], lw=1.5, linestyle="--")
    axes[0].set_xlabel("Predicted", fontsize=10)
    axes[0].set_ylabel("Residual",  fontsize=10)
    axes[0].set_title("Residual vs Predicted", fontsize=11, fontweight="bold")
    axes[0].spines[["top","right"]].set_visible(False)
    axes[1].hist(res, bins=40, color=CLR["teal"], edgecolor="white")
    axes[1].set_xlabel("Residual", fontsize=10)
    axes[1].set_ylabel("Count",    fontsize=10)
    axes[1].set_title("Residual Distribution", fontsize=11, fontweight="bold")
    axes[1].spines[["top","right"]].set_visible(False)
    fig.suptitle(title, fontsize=13, fontweight="bold", y=1.02)
    fig.tight_layout()
    return fig

# ── D14: actual vs predicted scatter ─────────────────────────────────────────
def plot_actual_pred(y_true, y_pred, title="Actual vs Predicted"):
    fig, ax = plt.subplots(figsize=(5, 4))
    lo, hi  = min(y_true.min(), y_pred.min()), max(y_true.max(), y_pred.max())
    ax.plot([lo, hi], [lo, hi], "--", color=CLR["danger"], lw=1.5,
            label="Perfect fit")
    ax.scatter(y_true, y_pred, alpha=.35, color=CLR["primary"],
               edgecolors="white", s=15)
    ax.set_xlabel("Actual",    fontsize=10)
    ax.set_ylabel("Predicted", fontsize=10)
    ax.set_title(title, fontsize=12, fontweight="bold", pad=10)
    ax.legend(fontsize=9)
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return fig

# ── D15: update stage2 insights (feeds Tab 6 of EDA dashboard) ───────────────
def update_stage2_insights():
    lines = ["=" * 60,
             "📊  STAGE 2 — ML MODEL RESULTS",
             f"Generated: {datetime.datetime.now():%Y-%m-%d  %H:%M}",
             "=" * 60]
    if S["reg_results"]:
        best = max(S["reg_results"], key=lambda k: S["reg_results"][k]["r2"])
        r    = S["reg_results"][best]
        lines += ["", "── REGRESSION ──",
                  f"Best Model  : {best}",
                  f"R²          : {r['r2']:.4f}",
                  f"MAE         : {r['mae']:,.0f}",
                  f"RMSE        : {r['rmse']:,.0f}",
                  f"CV R² mean  : {r['cv_mean']:.4f} ± {r['cv_std']:.4f}"]
    if S["cls_results"]:
        best = max(S["cls_results"], key=lambda k: S["cls_results"][k]["acc"])
        r    = S["cls_results"][best]
        lines += ["", "── CLASSIFICATION ──",
                  f"Best Model  : {best}",
                  f"Accuracy    : {r['acc']*100:.2f}%",
                  f"F1 Score    : {r['f1']:.4f}",
                  f"CV Acc mean : {r['cv_mean']*100:.2f}% ± {r['cv_std']*100:.2f}%"]
    S["stage2_insights"] = "\n".join(lines)

# ── D16: fig → bytes ─────────────────────────────────────────────────────────
def fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    return buf.read()

# ── D17: colour for R² ───────────────────────────────────────────────────────
def r2_colour(v: float) -> str:
    if v >= .85: return "green"
    if v >= .70: return ""
    return "amber"

def acc_colour(v: float) -> str:
    if v >= .85: return "green"
    if v >= .70: return ""
    return "amber"


# ─────────────────────────────────────────────────────────────────────────────
# E  SIDEBAR — DATA LOADER
# ─────────────────────────────────────────────────────────────────────────────
 



with st.sidebar:
    st.image(str(LOGO), width=70)
    st.markdown("---")

with st.sidebar:
    st.markdown("## 🤖 ML Models Engine")
    st.markdown("**Stage 2 — Training & Prediction**")
    st.divider()

    # ── Data source ──────────────────────────────────────────────────────────
    st.markdown("### 📂 Data Source")
    data_src = st.radio("", ["From EDA Dashboard (session)",
                              "Upload CSV file"], label_visibility="collapsed")

    df_global = None
    if data_src == "From EDA Dashboard (session)":
        df_global = load_data()
        if df_global is not None:
            st.success(f"✅ Loaded — {len(df_global):,} rows")
        else:
            st.warning("No data in session.\nRun ML_EDA_Dashboard first or upload a file.")
    else:
        up = st.file_uploader("Upload CSV", type=["csv"],
                               label_visibility="collapsed")
        if up:
            df_global = pd.read_csv(up)
            S["df_clean"] = df_global
            st.success(f"✅ {len(df_global):,} rows loaded")

    # ── Target / features ────────────────────────────────────────────────────
    if df_global is not None:
        st.divider()
        st.markdown("### ⚙️ Configuration")
        num_cols = df_global.select_dtypes(include=np.number).columns.tolist()
        target   = st.selectbox("Target column (regression)",
                                 num_cols,
                                 index=num_cols.index("MonthlyIncome")
                                 if "MonthlyIncome" in num_cols else 0)
        S["target_col"] = target
        feat_opts = [c for c in num_cols if c != target]
        sel_feats = st.multiselect("Feature columns", feat_opts,
                                   default=feat_opts[:12] if len(feat_opts) >= 12
                                   else feat_opts)
        if sel_feats:
            S["feat_names"] = sel_feats

        st.divider()
        st.markdown("### 🗂️ Price Categories")
        bins_raw  = st.text_input("Bin edges (comma-sep)",
                                  "0, 300000, 600000, 1000000, 9999999999")
        labs_raw  = st.text_input("Labels (comma-sep)",
                                  "Low, Medium, High, Luxury")
        try:
            S["price_bins"]   = [float(x) for x in bins_raw.split(",")]
            S["price_labels"] = [x.strip() for x in labs_raw.split(",")]
        except Exception:
            st.error("Invalid bins / labels")

    st.divider()
    st.markdown("### 💾 Model Persistence")
    save_dir = st.text_input("Save directory", "saved_models")

    st.divider()
    st.caption("ML Models Engine v2.0 · May 2025")


# ─────────────────────────────────────────────────────────────────────────────
# F  MAIN — 4 TABS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<div style="background:linear-gradient(135deg,#1E3A8A,#0891B2);
            border-radius:14px;padding:22px 28px;margin-bottom:20px;
            box-shadow:0 4px 20px rgba(0,0,0,.15);">
  <h1 style="color:white;margin:0;font-size:26px;">🤖 ML Models Engine — Stage 2</h1>
  <p style="color:#BAE6FD;margin:6px 0 0;font-size:14px;">
      Regression · Classification · Reports · Prediction
  </p>
</div>
""", unsafe_allow_html=True)

tab9, tab10, tab11, tab12,tab13 = st.tabs([
    "📉 Tab 9 · Regression Models",
    "🎯 Tab 10 · Classification Models",
    "📊 Tab 11 · Comparison & Reports",
    "🔮 Tab 12 · Predict New Data",
    "💡 Tab 13 · FINAL INSIGHTS & RECOMMENDATIONS",    # NEW
    
])

# guard — no data
def _no_data_msg():
    st.markdown('<div class="warning-box">⚠️ No data loaded. '
                'Use the sidebar to load a dataset.</div>',
                unsafe_allow_html=True)


# =============================================================================
# TAB 9 — REGRESSION MODELS
# =============================================================================
with tab9:
    st.markdown("## 📉 Regression Models")
    if df_global is None:
        _no_data_msg()
        st.stop()

    # ── 9.1 Define the 6 models ─────────────────────────────────────────────
    REG_MODELS = {
        "Linear Regression"          : LinearRegression(),
        "Ridge Regression"           : Ridge(alpha=1.0),
        "Lasso Regression"           : Lasso(alpha=0.1, max_iter=5000),
        "Decision Tree"              : DecisionTreeRegressor(max_depth=8,
                                           random_state=42),
        "Random Forest"              : RandomForestRegressor(n_estimators=150,
                                           max_depth=12, random_state=42,
                                           n_jobs=-1),
        "Gradient Boosting"          : GradientBoostingRegressor(
                                           n_estimators=150, learning_rate=.08,
                                           max_depth=5, random_state=42),
    }

    # ── 9.2 Prepare data ────────────────────────────────────────────────────
    if not S["data_prepared_r"]:
        prepare_regression(df_global)

    # ── 9.3 Controls ────────────────────────────────────────────────────────
    section("⚙️ Training Configuration")
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        n_cv_r = st.slider("CV folds", 3, 10, 5, key="cv_r")
    with col_b:
        sel_models_r = st.multiselect("Models to train",
                                       list(REG_MODELS.keys()),
                                       default=list(REG_MODELS.keys()),
                                       key="sel_r")
    with col_c:
        use_parallel_r = st.checkbox("⚡ Parallel training", value=True,
                                      key="par_r")
        if use_parallel_r:
            max_cores = os.cpu_count()          # 16 on your machine
            n_jobs_r  = st.slider(
                "⚡ Cores to use",
                min_value=1,
                max_value=max_cores,
                value=st.session_state.get("njobs_r_val", max(1, max_cores - 1)),
                step=1,
                key="njobs_r",
                help=f"Your CPU has {max_cores} cores. "
                     f"Recommended: leave at least 1 free for system stability."
            )
            st.session_state["njobs_r_val"] = n_jobs_r
           # Safe mode indicator
            if n_jobs_r == max_cores:
                st.warning("⚠️ Using ALL cores — system may slow down")
            elif n_jobs_r >= max_cores - 2:
                st.info(f"✅ Safe mode — {max_cores - n_jobs_r} core(s) free")
            else:
                st.success(f"✅ Conservative — {max_cores - n_jobs_r} cores free")
        else:
            n_jobs_r = 1
            st.session_state["njobs_r_val"] = 1

    # ── 9.4 Train button ─────────────────────────────────────────────────────
    col_btn, col_clr = st.columns([2, 1])
    with col_btn:
        run_r = st.button("🚀  Train Regression Models", type="primary",
                          use_container_width=True, key="train_r_btn")
    with col_clr:
        if st.button("🗑️ Clear Results", use_container_width=True,
                     key="clr_r_btn"):
            S["reg_results"] = {}
            S["reg_models"]  = {}
            st.rerun()
#--------------------------------------------------------------------
# ── CPU sidebar display ──────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("### ⚙️ Training Engine")
        cpu_info = get_cpu_info(use_parallel_r, n_jobs_r)
        st.metric("Total CPU Cores", cpu_info["total"])
        st.metric("Cores Selected",  n_jobs_r,
                  delta="Parallel 🟢" if use_parallel_r else "Sequential 🔴")
        st.metric("CPU Load",        f"{cpu_info['percent']} %")
        if use_parallel_r:
            safe  = max(1, os.cpu_count() - 1)
            used_pct = round(n_jobs_r / os.cpu_count() * 100)
            st.progress(n_jobs_r / os.cpu_count(),
                        text=f"Using {used_pct}% of cores  "
                             f"({'⚠️ All cores!' if n_jobs_r == os.cpu_count() else '✅ Safe mode'})")
#-------------------------------------------------------------------
    if run_r and sel_models_r:
        prepare_regression(df_global)          # refresh split
        progress = st.progress(0, text="Training models …")
        #-------------------------------
        # ── In-tab CPU metrics ───────────────────────────────────────────────
        cpu_info = get_cpu_info(use_parallel_r, n_jobs_r)
        c1, c2, c3 = st.columns(3)
        c1.metric("🖥️ Total Cores",   cpu_info["total"])
        c2.metric("⚡ Cores Selected", n_jobs_r,
                  delta="Parallel" if use_parallel_r else "Sequential")
        c3.metric("📊 CPU Load",      f"{cpu_info['percent']} %")
        #-------------------------------
        status   = st.empty()
        results  = {}
#---------------------------------------------------------------------
        if use_parallel_r:
            # ── Parallel via ThreadPoolExecutor ──────────────────────────────
            args_list = [(n, REG_MODELS[n],
                          S["X_train_r"], S["X_test_r"],
                          S["y_train_r"].values, S["y_test_r"].values)
                         for n in sel_models_r]
            with concurrent.futures.ThreadPoolExecutor(max_workers=n_jobs_r) as exe:
                futs = {exe.submit(_train_one_reg, a): a[0] for a in args_list}
                done = 0
                for fut in concurrent.futures.as_completed(futs):
                    name, res, err = fut.result()
                    done += 1
                    progress.progress(done / len(sel_models_r),
                                      text=f"✅ {name}")
                    if err:
                        st.error(f"{name}: {err}")
                    else:
                        results[name] = res
        else:
            # ── Sequential ───────────────────────────────────────────────────
            for i, name in enumerate(sel_models_r):
                status.info(f"Training {name} …")
                _, res, err = _train_one_reg((name, REG_MODELS[name],
                                              S["X_train_r"], S["X_test_r"],
                                              S["y_train_r"].values,
                                              S["y_test_r"].values))
                progress.progress((i + 1) / len(sel_models_r),
                                   text=f"✅ {name}")
                if err:
                    st.error(f"{name}: {err}")
                else:
                    results[name] = res

        S["reg_results"].update(results)
        S["reg_models"].update({n: r["model"] for n, r in results.items()})
        best = max(results, key=lambda k: results[k]["r2"])
        S["best_reg_name"] = best
        update_stage2_insights()
        status.empty()
        progress.progress(1.0, text="✅ All models trained!")
        st.success(f"🏆 Best model: **{best}**  —  R² = {results[best]['r2']:.4f}")
#---------------------------------------------------------------------------

    # ── 9.5 Results dashboard ───────────────────────────────────────────────
    if S["reg_results"]:
        section("📊 Results Overview")

        # Summary table
        rows = []
        for name, r in S["reg_results"].items():
            rows.append({"Model": name,
                          "R²"   : f"{r['r2']:.4f}",
                          "MAE"  : f"{r['mae']:,.0f}",
                          "RMSE" : f"{r['rmse']:,.0f}",
                          "CV R² (mean±std)": f"{r['cv_mean']:.4f} ± {r['cv_std']:.4f}"})
        df_sum = pd.DataFrame(rows).sort_values("R²", ascending=False)
        st.dataframe(df_sum, use_container_width=True, hide_index=True)

        # Best model metric cards
        best_n = max(S["reg_results"], key=lambda k: S["reg_results"][k]["r2"])
        best_r = S["reg_results"][best_n]
        st.markdown(f"**🏆 Best Model: {best_n}**")
        c1, c2, c3, c4 = st.columns(4)
        c1.markdown(metric_card("R² Score",
                                 f"{best_r['r2']:.4f}",
                                 "Higher is better",
                                 r2_colour(best_r["r2"])),
                    unsafe_allow_html=True)
        c2.markdown(metric_card("MAE",
                                 f"{best_r['mae']:,.0f}",
                                 "Mean Absolute Error", "teal"),
                    unsafe_allow_html=True)
        c3.markdown(metric_card("RMSE",
                                 f"{best_r['rmse']:,.0f}",
                                 "Root Mean Sq Error", "amber"),
                    unsafe_allow_html=True)
        c4.markdown(metric_card("CV R²",
                                 f"{best_r['cv_mean']:.4f}",
                                 f"±{best_r['cv_std']:.4f}  ·  {n_cv_r}-fold"),
                    unsafe_allow_html=True)

        # ── 9.6 Per-model deep dive ─────────────────────────────────────────
        section("🔍 Per-Model Analysis")
        chosen_r = st.selectbox("Select model for deep dive",
                                 list(S["reg_results"].keys()), key="dive_r")
        r_dive = S["reg_results"][chosen_r]
        model  = r_dive["model"]
        y_pr   = r_dive["y_pred"]
        y_te   = S["y_test_r"].values

        c1, c2 = st.columns(2)
        with c1:
            st.pyplot(plot_actual_pred(y_te, y_pr,
                                        f"{chosen_r} — Actual vs Predicted"))
        with c2:
            st.pyplot(plot_residuals(y_te, y_pr, f"{chosen_r} — Residuals"))

        c3, c4 = st.columns(2)
        with c3:
            st.pyplot(plot_feature_importance(model, S["feat_names"],
                                               f"{chosen_r} — Feature Importance"))
        with c4:
            if st.checkbox("Show learning curve", key="lc_r"):
                with st.spinner("Computing learning curve …"):
                    fig_lc = plot_learning_curve(
                        model, S["X_train_r"], S["y_train_r"].values,
                        "reg", f"{chosen_r} — Learning Curve")
                    st.pyplot(fig_lc)

        # ── 9.7 R² bar comparison ───────────────────────────────────────────
        section("📈 R² Comparison — All Models")
        names_r = list(S["reg_results"].keys())
        vals_r  = [S["reg_results"][n]["r2"] for n in names_r]
        colours_r = [CLR["success"] if v >= .85
                     else (CLR["primary"] if v >= .70 else CLR["warning"])
                     for v in vals_r]
        fig_cmp, ax_cmp = plt.subplots(figsize=(9, 3.5))
        bars = ax_cmp.barh(names_r, vals_r, color=colours_r,
                            edgecolor="white", height=.55)
        ax_cmp.bar_label(bars, fmt="%.4f", padding=4, fontsize=10)
        ax_cmp.set_xlim(0, 1)
        ax_cmp.axvline(.7, color=CLR["warning"], lw=1.2, ls="--",
                        label="70% threshold")
        ax_cmp.axvline(.85, color=CLR["success"], lw=1.2, ls="--",
                         label="85% threshold")
        ax_cmp.set_title("R² by Model", fontsize=13, fontweight="bold")
        ax_cmp.legend(fontsize=9)
        ax_cmp.spines[["top","right"]].set_visible(False)
        fig_cmp.tight_layout()
        st.pyplot(fig_cmp)

        # ── 9.8 Save model ──────────────────────────────────────────────────
        section("💾 Save Model")
        col_sv1, col_sv2 = st.columns(2)
        with col_sv1:
            model_to_save = st.selectbox("Model to save",
                                          list(S["reg_models"].keys()),
                                          key="save_r_sel")
        with col_sv2:
            if st.button("💾 Save to disk", key="save_r_btn"):
                os.makedirs(save_dir, exist_ok=True)
                path = os.path.join(save_dir, f"{model_to_save.replace(' ','_')}_reg.pkl")
                joblib.dump({"model": S["reg_models"][model_to_save],
                              "scaler": S["scaler_r"],
                              "features": S["feat_names"],
                              "target": S["target_col"],
                              "metrics": S["reg_results"][model_to_save]},
                             path)
                st.success(f"Saved → `{path}`")

        # In-memory download
        buf = io.BytesIO()
        joblib.dump({"model": S["reg_models"][list(S["reg_models"].keys())[0]],
                      "scaler": S["scaler_r"],
                      "features": S["feat_names"]}, buf)
        st.download_button("⬇️ Download best model (pkl)",
                            buf.getvalue(),
                            file_name=f"{best_n.replace(' ','_')}_reg.pkl",
                            mime="application/octet-stream",
                            key="dl_reg")


# =============================================================================
# TAB 10 — CLASSIFICATION MODELS
# =============================================================================
with tab10:
    st.markdown("## 🎯 Classification Models")
    if df_global is None:
        _no_data_msg()
        st.stop()

    # ── 10.1 Define 6 classifiers ────────────────────────────────────────────
    CLS_MODELS = {
        "Logistic Regression"  : LogisticRegression(max_iter=1000, class_weight="balanced",
                                      random_state=42),
        "K-Nearest Neighbors"  : KNeighborsClassifier(n_neighbors=7),
        "Decision Tree"        : DecisionTreeClassifier(max_depth=8, class_weight="balanced",
                                      random_state=42),
        "Random Forest"        : RandomForestClassifier(n_estimators=150, class_weight="balanced",
                                      max_depth=12, random_state=42,
                                      n_jobs=-1),
        "Gradient Boosting"    : GradientBoostingClassifier(n_estimators=150,
                                      learning_rate=.08, max_depth=5,
                                      random_state=42),
        "SVM"                  : SVC(kernel="rbf", probability=True, class_weight="balanced",
                                      random_state=42),
    }

    # ── 10.2 Prepare split ───────────────────────────────────────────────────
    if not S["data_prepared_c"]:
        prepare_classification(df_global)

    # ── 10.3 Controls ────────────────────────────────────────────────────────
    section("⚙️ Training Configuration")
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        n_cv_c = st.slider("CV folds", 3, 10, 5, key="cv_c")
    with col_b:
        sel_models_c = st.multiselect("Models to train",
                                       list(CLS_MODELS.keys()),
                                       default=list(CLS_MODELS.keys()),
                                       key="sel_c")
    with col_c:
        use_parallel_c = st.checkbox("⚡ Parallel training", value=True,
                                      key="par_c")

    col_btn2, col_clr2 = st.columns([2, 1])
    with col_btn2:
        run_c = st.button("🚀  Train Classification Models", type="primary",
                           use_container_width=True, key="train_c_btn")
    with col_clr2:
        if st.button("🗑️ Clear Results", use_container_width=True,
                     key="clr_c_btn"):
            S["cls_results"] = {}
            S["cls_models"]  = {}
            st.rerun()
#-------------------------------------------------
# ── CPU sidebar display ──────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("### ⚙️ Training Engine")
        cpu_info = get_cpu_info(use_parallel_r, n_jobs_r)
        st.metric("Total CPU Cores",  cpu_info["total"])
        st.metric("Cores Used",       cpu_info["used"],
                  delta="Parallel 🟢" if use_parallel_r else "Sequential 🔴")
        st.metric("CPU Load",         f"{cpu_info['percent']} %")
#-------------------------------------------------
    if run_c and sel_models_c:
        prepare_classification(df_global)
        progress2 = st.progress(0, text="Training models …")
#-----------------------------------------
# ── In-tab CPU metrics ───────────────────────────────────────────────
        cpu_info = get_cpu_info(use_parallel_r, n_jobs_r)
        c1, c2, c3 = st.columns(3)
        c1.metric("🖥️ Total Cores",  cpu_info["total"])
        c2.metric("⚡ Cores in Use", cpu_info["used"],
                  delta="Parallel" if use_parallel_r else "Sequential")
        c3.metric("📊 CPU Load",     f"{cpu_info['percent']} %")
#----------------------------------------                
        status2   = st.empty()
        results2  = {}

        if use_parallel_c:
            args_list2 = [(n, CLS_MODELS[n],
                           S["X_train_c"], S["X_test_c"],
                           S["y_train_c"],  S["y_test_c"])
                          for n in sel_models_c]
            with concurrent.futures.ThreadPoolExecutor() as exe:
                futs2 = {exe.submit(_train_one_cls, a): a[0] for a in args_list2}
                done2 = 0
                for fut in concurrent.futures.as_completed(futs2):
                    name, res, err = fut.result()
                    done2 += 1
                    progress2.progress(done2 / len(sel_models_c),
                                       text=f"✅ {name}")
                    if err:
                        st.error(f"{name}: {err}")
                    else:
                        results2[name] = res
        else:
            for i, name in enumerate(sel_models_c):
                status2.info(f"Training {name} …")
                _, res, err = _train_one_cls((name, CLS_MODELS[name],
                                              S["X_train_c"], S["X_test_c"],
                                              S["y_train_c"],  S["y_test_c"]))
                progress2.progress((i + 1) / len(sel_models_c),
                                    text=f"✅ {name}")
                if err:
                    st.error(f"{name}: {err}")
                else:
                    results2[name] = res

        S["cls_results"].update(results2)
        S["cls_models"].update({n: r["model"] for n, r in results2.items()})
        best_c = max(results2, key=lambda k: results2[k]["acc"])
        S["best_cls_name"] = best_c
        update_stage2_insights()
        status2.empty()
        progress2.progress(1.0, text="✅ All classifiers trained!")
        st.success(f"🏆 Best classifier: **{best_c}**  —  Acc = {results2[best_c]['acc']*100:.2f}%")

    # ── 10.4 Results ─────────────────────────────────────────────────────────
    if S["cls_results"]:
        section("📊 Results Overview")
        rows2 = []
        for name, r in S["cls_results"].items():
            rows2.append({"Model"    : name,
                           "Accuracy" : f"{r['acc']*100:.2f}%",
                           "F1"       : f"{r['f1']:.4f}",
                           "Precision": f"{r['precision']:.4f}",
                           "Recall"   : f"{r['recall']:.4f}",
                           "CV Acc (mean±std)":
                               f"{r['cv_mean']*100:.2f}% ± {r['cv_std']*100:.2f}%"})
        df_sum2 = pd.DataFrame(rows2).sort_values("Accuracy", ascending=False)
        st.dataframe(df_sum2, use_container_width=True, hide_index=True)

        best_cn = max(S["cls_results"], key=lambda k: S["cls_results"][k]["acc"])
        best_cr = S["cls_results"][best_cn]
        st.markdown(f"**🏆 Best Classifier: {best_cn}**")
        c1, c2, c3, c4 = st.columns(4)
        c1.markdown(metric_card("Accuracy",
                                 f"{best_cr['acc']*100:.2f}%",
                                 "Test set", acc_colour(best_cr["acc"])),
                    unsafe_allow_html=True)
        c2.markdown(metric_card("F1 Score",
                                 f"{best_cr['f1']:.4f}",
                                 "Weighted average", "teal"),
                    unsafe_allow_html=True)
        c3.markdown(metric_card("Precision",
                                 f"{best_cr['precision']:.4f}",
                                 "Weighted", "purple"),
                    unsafe_allow_html=True)
        c4.markdown(metric_card("CV Accuracy",
                                 f"{best_cr['cv_mean']*100:.2f}%",
                                 f"±{best_cr['cv_std']*100:.2f}%"),
                    unsafe_allow_html=True)

        # ── 10.5 Per-classifier deep dive ────────────────────────────────────
        section("🔍 Per-Classifier Analysis")
        chosen_c = st.selectbox("Select classifier",
                                  list(S["cls_results"].keys()), key="dive_c")
        r_cdive  = S["cls_results"][chosen_c]
        le_inv   = [str(c) for c in S["le"].classes_]  # must be strings for classification_report

        c1, c2 = st.columns(2)
        with c1:
            st.pyplot(plot_cm(r_cdive["cm"], le_inv,
                               f"{chosen_c} — Confusion Matrix"))
        with c2:
            st.pyplot(plot_feature_importance(r_cdive["model"],
                                               S["feat_names"],
                                               f"{chosen_c} — Feature Importance",
                                               CLR["purple"]))
        # Classification report
        with st.expander("📋 Full Classification Report"):
            cr_txt = classification_report(
                S["y_test_c"], r_cdive["y_pred"],
                target_names=le_inv, zero_division=0)
            st.code(cr_txt, language="text")

        # ROC — only if model supports predict_proba and binary / OvR
        if "y_prob" in r_cdive and len(le_inv) == 2:
            section("📈 ROC Curve")
            fpr, tpr, _ = roc_curve(S["y_test_c"], r_cdive["y_prob"][:, 1])
            auc_val      = roc_auc_score(S["y_test_c"], r_cdive["y_prob"][:, 1])
            fig_roc, ax_roc = plt.subplots(figsize=(5, 4))
            ax_roc.plot(fpr, tpr, color=CLR["primary"], lw=2,
                        label=f"AUC = {auc_val:.3f}")
            ax_roc.plot([0, 1], [0, 1], "--", color="#94A3B8")
            ax_roc.set_xlabel("FPR"); ax_roc.set_ylabel("TPR")
            ax_roc.set_title("ROC Curve", fontweight="bold")
            ax_roc.legend(); ax_roc.spines[["top","right"]].set_visible(False)
            fig_roc.tight_layout()
            st.pyplot(fig_roc)

        if st.checkbox("Show learning curve (classifier)", key="lc_c"):
            with st.spinner("Computing …"):
                fig_lcc = plot_learning_curve(
                    r_cdive["model"],
                    S["X_train_c"], S["y_train_c"],
                    "cls", f"{chosen_c} — Learning Curve")
                st.pyplot(fig_lcc)

        # ── 10.6 Accuracy bar comparison ─────────────────────────────────────
        section("📊 Accuracy Comparison — All Classifiers")
        names_c = list(S["cls_results"].keys())
        vals_c  = [S["cls_results"][n]["acc"] for n in names_c]
        colours_c = [CLR["success"] if v >= .85
                     else (CLR["primary"] if v >= .70 else CLR["warning"])
                     for v in vals_c]
        fig_cmp2, ax_cmp2 = plt.subplots(figsize=(9, 3.5))
        bars2 = ax_cmp2.barh(names_c, vals_c, color=colours_c,
                              edgecolor="white", height=.55)
        ax_cmp2.bar_label(bars2, fmt="%.3f", padding=4, fontsize=10)
        ax_cmp2.set_xlim(0, 1)
        ax_cmp2.axvline(.70, color=CLR["warning"], lw=1.2, ls="--")
        ax_cmp2.axvline(.85, color=CLR["success"], lw=1.2, ls="--")
        ax_cmp2.set_title("Accuracy by Classifier", fontsize=13,
                           fontweight="bold")
        ax_cmp2.spines[["top","right"]].set_visible(False)
        fig_cmp2.tight_layout()
        st.pyplot(fig_cmp2)

        # ── 10.7 Save ─────────────────────────────────────────────────────────
        section("💾 Save Classifier")
        col_sv3, col_sv4 = st.columns(2)
        with col_sv3:
            model_to_save_c = st.selectbox("Classifier to save",
                                            list(S["cls_models"].keys()),
                                            key="save_c_sel")
        with col_sv4:
            if st.button("💾 Save to disk", key="save_c_btn"):
                os.makedirs(save_dir, exist_ok=True)
                path_c = os.path.join(save_dir,
                             f"{model_to_save_c.replace(' ','_')}_cls.pkl")
                joblib.dump({"model"   : S["cls_models"][model_to_save_c],
                              "scaler"  : S["scaler_c"],
                              "le"      : S["le"],
                              "features": S["feat_names"],
                              "metrics" : S["cls_results"][model_to_save_c]},
                             path_c)
                st.success(f"Saved → `{path_c}`")

        buf_c = io.BytesIO()
        joblib.dump({"model"   : S["cls_models"][best_cn],
                      "scaler"  : S["scaler_c"],
                      "le"      : S["le"],
                      "features": S["feat_names"]}, buf_c)
        st.download_button("⬇️ Download best classifier (pkl)",
                            buf_c.getvalue(),
                            file_name=f"{best_cn.replace(' ','_')}_cls.pkl",
                            mime="application/octet-stream",
                            key="dl_cls")


# =============================================================================
# TAB 11 — COMPARISON & REPORTS
# =============================================================================
with tab11:
    st.markdown("## 📊 Model Comparison & Reports")

    have_r = bool(S["reg_results"])
    have_c = bool(S["cls_results"])

    if not have_r and not have_c:
        st.markdown('<div class="warning-box">⚠️ Train models in Tab 9 and/or '
                    'Tab 10 first.</div>', unsafe_allow_html=True)
    else:
        # ── 11.1 Side-by-side summary ────────────────────────────────────────
        section("🏆 Best Models at a Glance")
        col_r, col_c = st.columns(2)

        if have_r:
            best_rn = max(S["reg_results"], key=lambda k: S["reg_results"][k]["r2"])
            br      = S["reg_results"][best_rn]
            with col_r:
                st.markdown(f"""
                <div class="pred-card">
                  <div class="pred-label">🏅 Best Regressor</div>
                  <div style="font-size:20px;font-weight:700;color:#1E293B;">
                      {best_rn}</div>
                  <hr style="border-color:#E2E8F0;margin:12px 0">
                  <div style="display:flex;justify-content:space-around;">
                    <div><div class="metric-label">R²</div>
                         <div style="font-size:22px;font-weight:700;
                              color:#059669;">{br['r2']:.4f}</div></div>
                    <div><div class="metric-label">MAE</div>
                         <div style="font-size:22px;font-weight:700;
                              color:#2563EB;">{br['mae']:,.0f}</div></div>
                    <div><div class="metric-label">RMSE</div>
                         <div style="font-size:22px;font-weight:700;
                              color:#D97706;">{br['rmse']:,.0f}</div></div>
                  </div>
                </div>""", unsafe_allow_html=True)

        if have_c:
            best_cn2 = max(S["cls_results"], key=lambda k: S["cls_results"][k]["acc"])
            bc       = S["cls_results"][best_cn2]
            with col_c:
                st.markdown(f"""
                <div class="pred-card">
                  <div class="pred-label">🏅 Best Classifier</div>
                  <div style="font-size:20px;font-weight:700;color:#1E293B;">
                      {best_cn2}</div>
                  <hr style="border-color:#E2E8F0;margin:12px 0">
                  <div style="display:flex;justify-content:space-around;">
                    <div><div class="metric-label">Accuracy</div>
                         <div style="font-size:22px;font-weight:700;
                              color:#059669;">{bc['acc']*100:.2f}%</div></div>
                    <div><div class="metric-label">F1</div>
                         <div style="font-size:22px;font-weight:700;
                              color:#2563EB;">{bc['f1']:.4f}</div></div>
                    <div><div class="metric-label">CV Acc</div>
                         <div style="font-size:22px;font-weight:700;
                              color:#7C3AED;">{bc['cv_mean']*100:.2f}%</div></div>
                  </div>
                </div>""", unsafe_allow_html=True)

        # ── 11.2 Combined leaderboard chart ──────────────────────────────────
        if have_r and have_c:
            section("📈 Combined Performance Chart")
            fig_dual, (ax_l, ax_r) = plt.subplots(1, 2, figsize=(13, 4))

            # Regression — R²
            rns  = list(S["reg_results"].keys())
            rvs  = [S["reg_results"][n]["r2"] for n in rns]
            rcls = [CLR["success"] if v >= .85
                    else (CLR["primary"] if v >= .70 else CLR["warning"])
                    for v in rvs]
            b1   = ax_l.barh(rns, rvs, color=rcls, edgecolor="white",
                              height=.55)
            ax_l.bar_label(b1, fmt="%.3f", padding=3, fontsize=9)
            ax_l.set_xlim(0, 1)
            ax_l.set_title("R² — Regression Models", fontweight="bold")
            ax_l.spines[["top","right"]].set_visible(False)

            # Classification — Accuracy
            cns  = list(S["cls_results"].keys())
            cvs  = [S["cls_results"][n]["acc"] for n in cns]
            ccls = [CLR["success"] if v >= .85
                    else (CLR["primary"] if v >= .70 else CLR["warning"])
                    for v in cvs]
            b2   = ax_r.barh(cns, cvs, color=ccls, edgecolor="white",
                              height=.55)
            ax_r.bar_label(b2, fmt="%.3f", padding=3, fontsize=9)
            ax_r.set_xlim(0, 1)
            ax_r.set_title("Accuracy — Classification Models", fontweight="bold")
            ax_r.spines[["top","right"]].set_visible(False)

            fig_dual.tight_layout()
            st.pyplot(fig_dual)

        # ── 11.3 Detailed regression & classification tables ─────────────────
        if have_r:
            section("📋 Regression — Full Results")
            df_r_full = pd.DataFrame([
                {"Model": n, "R²": f"{r['r2']:.4f}",
                 "MAE": f"{r['mae']:,.0f}", "RMSE": f"{r['rmse']:,.0f}",
                 "CV R² mean": f"{r['cv_mean']:.4f}",
                 "CV R² std": f"{r['cv_std']:.4f}"}
                for n, r in S["reg_results"].items()
            ]).sort_values("R²", ascending=False)
            st.dataframe(df_r_full, use_container_width=True, hide_index=True)

        if have_c:
            section("📋 Classification — Full Results")
            df_c_full = pd.DataFrame([
                {"Model": n, "Accuracy": f"{r['acc']*100:.2f}%",
                 "F1": f"{r['f1']:.4f}",
                 "Precision": f"{r['precision']:.4f}",
                 "Recall": f"{r['recall']:.4f}",
                 "CV Acc mean": f"{r['cv_mean']*100:.2f}%",
                 "CV Acc std": f"{r['cv_std']*100:.2f}%"}
                for n, r in S["cls_results"].items()
            ]).sort_values("Accuracy", ascending=False)
            st.dataframe(df_c_full, use_container_width=True, hide_index=True)

        # ── 11.4 Stage 2 → Tab 8 feed preview ───────────────────────────────
        section("🔗 Stage 2 → EDA Dashboard (Tab 8 Feed)")
        st.markdown('<div class="info-box">The text block below is automatically '
                    'injected into <b>Tab 6 · Insights & Recommendations</b> of '
                    'the EDA Dashboard whenever models have been trained.</div>',
                    unsafe_allow_html=True)
        if S["stage2_insights"]:
            st.code(S["stage2_insights"], language="text")
        else:
            st.caption("No results yet — train models first.")

        # ── 11.5 Export Reports ───────────────────────────────────────────────
        section("📄 Export Reports")
        exp_col1, exp_col2 = st.columns(2)

        # ── PDF ──────────────────────────────────────────────────────────────
        with exp_col1:
            st.markdown("**PDF Report**")
            if st.button("📄 Generate PDF", key="gen_pdf",
                          use_container_width=True):
                if not REPORTLAB_OK:
                    st.error("Install reportlab: `pip install reportlab`")
                else:
                    buf_pdf = io.BytesIO()
                    doc     = SimpleDocTemplate(buf_pdf, pagesize=A4,
                                                topMargin=.7*inch,
                                                bottomMargin=.7*inch,
                                                leftMargin=.85*inch,
                                                rightMargin=.85*inch)
                    styles  = getSampleStyleSheet()
                    story   = []

                    # Title
                    title_style = ParagraphStyle("Title2",
                        parent=styles["Title"], fontSize=20,
                        textColor=colors.HexColor("#1E3A8A"), spaceAfter=6)
                    story.append(Paragraph("ML Models Engine — Report", title_style))
                    story.append(Paragraph(
                        f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M}",
                        styles["Normal"]))
                    story.append(HRFlowable(width="100%", thickness=1,
                                             color=colors.HexColor("#2563EB")))
                    story.append(Spacer(1, .2*inch))

                    h2_style = ParagraphStyle("H2",
                        parent=styles["Heading2"], fontSize=13,
                        textColor=colors.HexColor("#1E3A8A"), spaceBefore=10)

                    # Regression section
                    if have_r:
                        story.append(Paragraph("Regression Results", h2_style))
                        t_data = [["Model","R²","MAE","RMSE","CV R²"]]
                        for n, r in S["reg_results"].items():
                            t_data.append([n, f"{r['r2']:.4f}",
                                           f"{r['mae']:,.0f}",
                                           f"{r['rmse']:,.0f}",
                                           f"{r['cv_mean']:.4f}±{r['cv_std']:.4f}"])
                        tbl = Table(t_data,
                                    colWidths=[1.8*inch,.8*inch,.9*inch,
                                               .9*inch,1.3*inch])
                        tbl.setStyle(TableStyle([
                            ("BACKGROUND", (0,0), (-1,0),
                             colors.HexColor("#1E3A8A")),
                            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                            ("FONTSIZE",   (0,0), (-1,-1), 9),
                            ("GRID",       (0,0), (-1,-1), .3,
                             colors.HexColor("#CBD5E1")),
                            ("ROWBACKGROUNDS", (0,1), (-1,-1),
                             [colors.white, colors.HexColor("#F1F5F9")]),
                            ("ALIGN",      (1,1), (-1,-1), "CENTER"),
                        ]))
                        story.append(tbl)
                        story.append(Spacer(1, .15*inch))

                    # Classification section
                    if have_c:
                        story.append(Paragraph("Classification Results", h2_style))
                        t_data2 = [["Model","Accuracy","F1","Precision",
                                    "Recall","CV Acc"]]
                        for n, r in S["cls_results"].items():
                            t_data2.append([n,
                                            f"{r['acc']*100:.2f}%",
                                            f"{r['f1']:.4f}",
                                            f"{r['precision']:.4f}",
                                            f"{r['recall']:.4f}",
                                            f"{r['cv_mean']*100:.1f}%"])
                        tbl2 = Table(t_data2,
                                     colWidths=[1.7*inch,.9*inch,.7*inch,
                                                .8*inch,.7*inch,.9*inch])
                        tbl2.setStyle(TableStyle([
                            ("BACKGROUND", (0,0), (-1,0),
                             colors.HexColor("#059669")),
                            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                            ("FONTSIZE",   (0,0), (-1,-1), 9),
                            ("GRID",       (0,0), (-1,-1), .3,
                             colors.HexColor("#CBD5E1")),
                            ("ROWBACKGROUNDS", (0,1), (-1,-1),
                             [colors.white, colors.HexColor("#ECFDF5")]),
                            ("ALIGN",      (1,1), (-1,-1), "CENTER"),
                        ]))
                        story.append(tbl2)

                    # Stage2 feed
                    story.append(Spacer(1, .2*inch))
                    story.append(Paragraph("Stage 2 Feed (Tab 6)", h2_style))
                    story.append(Paragraph(
                        S["stage2_insights"].replace("\n","<br/>"),
                        styles["Code"]))

                    doc.build(story)
                    buf_pdf.seek(0)
                    st.download_button("⬇️ Download PDF",
                                        buf_pdf.getvalue(),
                                        file_name="ML_Report.pdf",
                                        mime="application/pdf",
                                        key="dl_pdf")

        # ── Word ──────────────────────────────────────────────────────────────
        with exp_col2:
            st.markdown("**Word Report (.docx)**")
            if st.button("📝 Generate Word", key="gen_word",
                          use_container_width=True):
                if not DOCX_OK:
                    st.error("Install python-docx: `pip install python-docx`")
                else:
                    wd_doc = Document()
                    wd_doc.core_properties.author = "ML Engine — Mohamed"

                    # Heading
                    h = wd_doc.add_heading("ML Models Engine — Report", 0)
                    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                    wd_doc.add_paragraph(
                        f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M}")
                    wd_doc.add_paragraph()

                    def add_table(doc, data_rows, headers):
                        tbl = doc.add_table(rows=1 + len(data_rows),
                                             cols=len(headers))
                        tbl.style = "Light Shading Accent 1"
                        for i, h_txt in enumerate(headers):
                            cell = tbl.rows[0].cells[i]
                            cell.text = h_txt
                            cell.paragraphs[0].runs[0].bold = True
                        for r_i, row in enumerate(data_rows, 1):
                            for c_i, val in enumerate(row):
                                tbl.rows[r_i].cells[c_i].text = str(val)

                    if have_r:
                        wd_doc.add_heading("Regression Results", 1)
                        reg_rows = [[n,
                                     f"{r['r2']:.4f}",
                                     f"{r['mae']:,.0f}",
                                     f"{r['rmse']:,.0f}",
                                     f"{r['cv_mean']:.4f}"]
                                    for n, r in S["reg_results"].items()]
                        add_table(wd_doc, reg_rows,
                                  ["Model","R²","MAE","RMSE","CV R²"])
                        wd_doc.add_paragraph()

                    if have_c:
                        wd_doc.add_heading("Classification Results", 1)
                        cls_rows = [[n,
                                     f"{r['acc']*100:.2f}%",
                                     f"{r['f1']:.4f}",
                                     f"{r['precision']:.4f}",
                                     f"{r['recall']:.4f}"]
                                    for n, r in S["cls_results"].items()]
                        add_table(wd_doc, cls_rows,
                                  ["Model","Accuracy","F1","Precision","Recall"])
                        wd_doc.add_paragraph()

                    wd_doc.add_heading("Stage 2 Feed (Tab 6)", 1)
                    wd_doc.add_paragraph(S["stage2_insights"])

                    buf_wd = io.BytesIO()
                    wd_doc.save(buf_wd)
                    buf_wd.seek(0)
                    st.download_button("⬇️ Download Word",
                                        buf_wd.getvalue(),
                                        file_name="ML_Report.docx",
                                        mime="application/vnd.openxmlformats-"
                                             "officedocument.wordprocessingml"
                                             ".document",
                                        key="dl_word")

        # ── CSV summary export ────────────────────────────────────────────────
        section("📥 Export Summary CSV")
        col_csv1, col_csv2 = st.columns(2)
        if have_r:
            df_r_csv = pd.DataFrame([
                {"Model":n,"R²":r["r2"],"MAE":r["mae"],"RMSE":r["rmse"],
                 "CV_R2_mean":r["cv_mean"],"CV_R2_std":r["cv_std"]}
                for n,r in S["reg_results"].items()])
            with col_csv1:
                st.download_button("⬇️ Regression CSV",
                                    df_r_csv.to_csv(index=False),
                                    "regression_results.csv", "text/csv",
                                    key="dl_r_csv")
        if have_c:
            df_c_csv = pd.DataFrame([
                {"Model":n,"Accuracy":r["acc"],"F1":r["f1"],
                 "Precision":r["precision"],"Recall":r["recall"],
                 "CV_Acc_mean":r["cv_mean"],"CV_Acc_std":r["cv_std"]}
                for n,r in S["cls_results"].items()])
            with col_csv2:
                st.download_button("⬇️ Classification CSV",
                                    df_c_csv.to_csv(index=False),
                                    "classification_results.csv","text/csv",
                                    key="dl_c_csv")


# =============================================================================
# TAB 12 — PREDICT NEW DATA
# =============================================================================
with tab12:
    st.markdown("## 🔮 Predict New Data")

    have_reg = bool(S["reg_models"])
    have_cls = bool(S["cls_models"])

    if not have_reg and not have_cls:
        st.markdown('<div class="warning-box">⚠️ No models trained yet. '
                    'Go to Tabs 9 & 10 to train models first.</div>',
                    unsafe_allow_html=True)
    else:
        # ── Always update stage2 feed silently ───────────────────────────────
        update_stage2_insights()

        # Model selectors
        section("⚙️ Model Selection")
        col_ms1, col_ms2 = st.columns(2)
        with col_ms1:
            if have_reg:
                sel_pred_r = st.selectbox(
                    "Regression model",
                    list(S["reg_models"].keys()), key="pred_r_sel")
        with col_ms2:
            if have_cls:
                sel_pred_c = st.selectbox(
                    "Classification model",
                    list(S["cls_models"].keys()), key="pred_c_sel")

        # ── 12.1 SINGLE ROW PREDICTION ───────────────────────────────────────
        section("🎯 Single-Row Prediction")
        feat_names = S["feat_names"]

        if df_global is not None and feat_names:
            ref_df = df_global[feat_names].describe()

            with st.expander("📝 Enter feature values", expanded=True):
                n_cols = 4
                input_vals = {}
                cols_inp   = st.columns(n_cols)
                for i, feat in enumerate(feat_names):
                    med = float(ref_df.loc["50%", feat]) \
                          if feat in ref_df.columns else 0.0
                    mn  = float(ref_df.loc["min", feat]) \
                          if feat in ref_df.columns else 0.0
                    mx  = float(ref_df.loc["max", feat]) \
                          if feat in ref_df.columns else 1e6
                    with cols_inp[i % n_cols]:
                        input_vals[feat] = st.number_input(
                            f"{feat}",
                            value=med,
                            help=f"Median: {med:,.1f}  |  "
                                 f"Range: [{mn:,.0f} – {mx:,.0f}]",
                            key=f"inp_{feat}",
                            format="%.2f")

            if st.button("🚀 Predict", type="primary", key="predict_btn"):
                row_df = pd.DataFrame([input_vals])

                # ── Regression prediction ────────────────────────────────────
                if have_reg:
                    row_scaled_r = S["scaler_r"].transform(
                        row_df[feat_names].values)
                    pred_price   = S["reg_models"][sel_pred_r].predict(
                        row_scaled_r)[0]
                    S["_last_pred_price"] = pred_price
                    S["_last_pred_model_r"] = sel_pred_r

                # ── Classification prediction ────────────────────────────────
                if have_cls:
                    row_scaled_c = S["scaler_c"].transform(
                        row_df[feat_names].values)
                    pred_cat_enc = S["cls_models"][sel_pred_c].predict(
                        row_scaled_c)[0]
                    pred_cat     = S["le"].inverse_transform([pred_cat_enc])[0]
                    S["_last_pred_cat"]   = pred_cat
                    S["_last_pred_model_c"] = sel_pred_c

                    # Confidence
                    if hasattr(S["cls_models"][sel_pred_c], "predict_proba"):
                        proba = S["cls_models"][sel_pred_c].predict_proba(
                            row_scaled_c)[0]
                        conf  = proba.max() * 100
                        S["_last_pred_conf"] = conf
                    else:
                        S["_last_pred_conf"] = None

            # ── Display results ───────────────────────────────────────────────
            if "_last_pred_price" in S or "_last_pred_cat" in S:
                st.markdown("---")
                res_c1, res_c2 = st.columns(2)

                if "_last_pred_price" in S and have_reg:
                    with res_c1:
                        pp = S["_last_pred_price"]
                        st.markdown(f"""
                        <div class="pred-card">
                          <div class="pred-label">📉 Regression Result
                              ({S['_last_pred_model_r']})</div>
                          <div class="pred-price">${pp:,.0f}</div>
                          <div class="pred-label">Predicted Price</div>
                        </div>""", unsafe_allow_html=True)

                if "_last_pred_cat" in S and have_cls:
                    with res_c2:
                        cat  = S["_last_pred_cat"]
                        conf = S.get("_last_pred_conf")
                        badge_cls, icon = CAT_STYLE.get(cat, ("badge-purple","⬜"))
                        st.markdown(f"""
                        <div class="pred-card">
                          <div class="pred-label">🎯 Classification Result
                              ({S['_last_pred_model_c']})</div>
                          <div class="pred-cat">{icon} {cat}</div>
                          <div style="margin-top:8px;">
                              <span class="badge {badge_cls}">
                                  Price Category</span>
                          </div>
                          {"<div class='pred-label' style='margin-top:10px;'>"
                           f"Confidence: <b>{conf:.1f}%</b></div>"
                           if conf is not None else ""}
                        </div>""", unsafe_allow_html=True)

                # Input summary
                with st.expander("📋 Input Summary"):
                    st.dataframe(pd.DataFrame([input_vals]).T.rename(
                        columns={0: "Value"}), use_container_width=True)

        # ── 12.2 BATCH PREDICTION ─────────────────────────────────────────────
        section("📂 Batch Prediction (CSV / Excel)")
        st.markdown('<div class="info-box">Upload a file with the same features '
                    'used during training. Missing columns will be filled with 0.'
                    '</div>', unsafe_allow_html=True)

        batch_file = st.file_uploader("Upload file for batch prediction",
                                       type=["csv","xlsx","xls"],
                                       key="batch_upload")
        if batch_file:
            try:
                if batch_file.name.endswith(".csv"):
                    df_batch = pd.read_csv(batch_file)
                else:
                    df_batch = pd.read_excel(batch_file)

                st.success(f"✅ {len(df_batch):,} rows loaded")

                # Align columns
                for feat in feat_names:
                    if feat not in df_batch.columns:
                        df_batch[feat] = 0

                X_batch = df_batch[feat_names].fillna(0)

                if st.button("🚀 Run Batch Prediction", type="primary",
                              key="batch_btn"):
                    with st.spinner("Running predictions …"):
                        df_out = df_batch.copy()

                        # Regression
                        if have_reg:
                            X_b_r = S["scaler_r"].transform(X_batch)
                            df_out["Predicted_Price"] = \
                                S["reg_models"][sel_pred_r].predict(X_b_r)

                        # Classification
                        if have_cls:
                            X_b_c = S["scaler_c"].transform(X_batch)
                            enc_cats = S["cls_models"][sel_pred_c].predict(X_b_c)
                            df_out["Predicted_Category"] = \
                                S["le"].inverse_transform(enc_cats)

                            if hasattr(S["cls_models"][sel_pred_c],
                                       "predict_proba"):
                                probas = S["cls_models"][sel_pred_c]\
                                             .predict_proba(X_b_c)
                                df_out["Confidence_%"] = \
                                    (probas.max(axis=1) * 100).round(2)

                        S["batch_results"] = df_out

                    st.success("✅ Batch prediction complete!")

            except Exception as e:
                st.error(f"Error reading file: {e}")

        # ── Display batch results ─────────────────────────────────────────────
        if S["batch_results"] is not None:
            df_out = S["batch_results"]
            section("📊 Batch Results")

            # Summary metrics
            mc1, mc2, mc3 = st.columns(3)
            mc1.markdown(metric_card("Total Rows", f"{len(df_out):,}",
                                      "Processed", "teal"),
                          unsafe_allow_html=True)
            if "Predicted_Price" in df_out.columns:
                mc2.markdown(metric_card("Avg Predicted Price",
                                          f"${df_out['Predicted_Price'].mean():,.0f}",
                                          f"Range: ${df_out['Predicted_Price'].min():,.0f}"
                                          f" – ${df_out['Predicted_Price'].max():,.0f}",
                                          "green"),
                              unsafe_allow_html=True)
            if "Predicted_Category" in df_out.columns:
                top_cat = df_out["Predicted_Category"].value_counts().idxmax()
                mc3.markdown(metric_card("Most Common Category",
                                          top_cat, "", "purple"),
                              unsafe_allow_html=True)

            # Category distribution pie
            if "Predicted_Category" in df_out.columns:
                fig_pie, ax_pie = plt.subplots(figsize=(5, 4))
                cat_counts = df_out["Predicted_Category"].value_counts()
                pie_colours = {"Low":"#059669","Medium":"#2563EB",
                               "High":"#D97706","Luxury":"#DC2626"}
                clrs = [pie_colours.get(c,"#7C3AED") for c in cat_counts.index]
                ax_pie.pie(cat_counts.values, labels=cat_counts.index,
                            autopct="%1.1f%%", colors=clrs,
                            startangle=90, pctdistance=.82)
                ax_pie.set_title("Category Distribution", fontweight="bold")
                fig_pie.tight_layout()
                col_pie, col_tbl = st.columns([1, 2])
                with col_pie:
                    st.pyplot(fig_pie)
                with col_tbl:
                    st.dataframe(df_out.head(500), use_container_width=True)
            else:
                st.dataframe(df_out.head(500), use_container_width=True)

            if len(df_out) > 500:
                st.caption(f"Showing first 500 of {len(df_out):,} rows.")

            # ── Export batch results ──────────────────────────────────────────
            section("📥 Export Batch Results")
            col_ex1, col_ex2 = st.columns(2)

            with col_ex1:
                csv_bytes = df_out.to_csv(index=False).encode()
                st.download_button("⬇️ Download CSV",
                                    csv_bytes, "batch_predictions.csv",
                                    "text/csv", key="dl_batch_csv")

            with col_ex2:
                buf_xl = io.BytesIO()
                with pd.ExcelWriter(buf_xl, engine="openpyxl") as writer:
                    df_out.to_excel(writer, sheet_name="Predictions",
                                     index=False)
                    # Summary sheet
                    if "Predicted_Price" in df_out.columns:
                        summary = pd.DataFrame({
                            "Metric": ["Count","Mean","Median","Min","Max","Std"],
                            "Value" : [len(df_out),
                                       df_out["Predicted_Price"].mean(),
                                       df_out["Predicted_Price"].median(),
                                       df_out["Predicted_Price"].min(),
                                       df_out["Predicted_Price"].max(),
                                       df_out["Predicted_Price"].std()]
                        })
                        summary.to_excel(writer, sheet_name="Summary",
                                          index=False)
                buf_xl.seek(0)
                st.download_button("⬇️ Download Excel (2 sheets)",
                                    buf_xl.getvalue(),
                                    "batch_predictions.xlsx",
                                    "application/vnd.openxmlformats-"
                                    "officedocument.spreadsheetml.sheet",
                                    key="dl_batch_xl")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 13 — FINAL INSIGHTS & RECOMMENDATIONS
# ─────────────────────────────────────────────────────────────────────────────
with tab13:
    st.markdown('<div style="background:linear-gradient(135deg,#1E3A8A,#0891B2);\
                border-radius:14px;padding:18px 24px;margin-bottom:20px;">\
                <h2 style="color:white;margin:0;font-size:22px;">📋 Final Insights & Recommendations</h2>\
                <p style="color:#BAE6FD;margin:4px 0 0;font-size:13px;">\
                Stage 1 EDA + Stage 2 ML Results — Combined Report</p></div>',
                unsafe_allow_html=True)

    if st.button("📋 Generate Final Report", type="primary",
                 key="btn_final_report"):

        update_stage2_insights()   # refresh Stage 2 results first

        lines = []
        lines.append("=" * 65)
        lines.append("   FINAL INSIGHTS & RECOMMENDATIONS REPORT")
        lines.append(f"   Generated : {datetime.datetime.now():%Y-%m-%d  %H:%M}")
        lines.append("=" * 65)

        # ── Section 1: Dataset Info ───────────────────────────────────────
        lines.append("")
        lines.append("1. DATASET SUMMARY")
        lines.append("─" * 45)
        if S.get("file_name"):
            lines.append(f"   File    : {S.get('file_name', 'N/A')}")
        if S.get("df_work") is not None:
            df_w = S["df_work"]
            lines.append(f"   Shape   : {df_w.shape[0]:,} rows × {df_w.shape[1]} columns")
        if S.get("target_col"):
            lines.append(f"   Target  : {S['target_col']}")
        if S.get("important_vars"):
            lines.append(f"   Features selected : {len(S['important_vars'])}")
            lines.append(f"   {', '.join(S['important_vars'][:8])}"
                         + (" …" if len(S["important_vars"]) > 8 else ""))

        # ── Section 2: EDA Insights from Stage 1 ─────────────────────────
        lines.append("")
        lines.append("2. EDA INSIGHTS  (Stage 1 Summary)")
        lines.append("─" * 45)
        stage1_text = S.get("insights_text", "")
        if stage1_text:
            # Extract only the key findings, skip header/footer lines
            for line in stage1_text.split("\n"):
                stripped = line.strip()
                if stripped and not stripped.startswith("=") \
                        and "ML Engine" not in stripped \
                        and "Report generated" not in stripped:
                    lines.append(f"   {stripped}")
        else:
            lines.append("   ⚠️  Run Tab 8 (Insights) in Stage 1 first")
            lines.append("       to include EDA findings here.")

        # ── Section 3: Regression Results ────────────────────────────────
        lines.append("")
        lines.append("3. REGRESSION MODEL RESULTS  (Stage 2)")
        lines.append("─" * 45)
        if S.get("reg_results"):
            sorted_reg = sorted(S["reg_results"].items(),
                                key=lambda x: x[1]["r2"], reverse=True)
            for rank, (name, r) in enumerate(sorted_reg, 1):
                lines.append(f"   #{rank}  {name}")
                lines.append(f"       R²      = {r['r2']:.4f}")
                lines.append(f"       MAE     = {r['mae']:,.0f}")
                lines.append(f"       RMSE    = {r['rmse']:,.0f}")
                lines.append(f"       CV R²   = {r['cv_mean']:.4f} ± {r['cv_std']:.4f}")
                lines.append("")
            best_r = sorted_reg[0][0]
            lines.append(f"   🏆 BEST REGRESSION MODEL : {best_r}")
            lines.append(f"      R² = {S['reg_results'][best_r]['r2']:.4f}")
        else:
            lines.append("   ⚠️  No regression results yet.")
            lines.append("       Train models in Tab 9 first.")

        # ── Section 4: Classification Results ────────────────────────────
        lines.append("")
        lines.append("4. CLASSIFICATION MODEL RESULTS  (Stage 2)")
        lines.append("─" * 45)
        if S.get("cls_results"):
            sorted_cls = sorted(S["cls_results"].items(),
                                key=lambda x: x[1]["acc"], reverse=True)
            for rank, (name, r) in enumerate(sorted_cls, 1):
                lines.append(f"   #{rank}  {name}")
                lines.append(f"       Accuracy = {r['acc']*100:.2f}%")
                lines.append(f"       F1 Score = {r['f1']:.4f}")
                lines.append(f"       CV Acc   = {r['cv_mean']*100:.2f}% ± "
                             f"{r['cv_std']*100:.2f}%")
                lines.append("")
            best_c = sorted_cls[0][0]
            lines.append(f"   🏆 BEST CLASSIFICATION MODEL : {best_c}")
            lines.append(f"      Accuracy = {S['cls_results'][best_c]['acc']*100:.2f}%")
        else:
            lines.append("   ⚠️  No classification results yet.")
            lines.append("       Train models in Tab 10 first.")

        # ── Section 5: Final Recommendations ─────────────────────────────
        lines.append("")
        lines.append("5. FINAL RECOMMENDATIONS")
        lines.append("─" * 45)
        lines.append("   EDA Stage:")
        lines.append("   • Verify missing value imputation (Tab 6)")
        lines.append("   • Check VIF multicollinearity (Tab 7)")
        lines.append("   • Review outlier treatment (Tab 3)")
        lines.append("")
        lines.append("   Modelling Stage:")
        if S.get("reg_results"):
            best_r = max(S["reg_results"], key=lambda k: S["reg_results"][k]["r2"])
            lines.append(f"   • Use {best_r} for regression tasks")
        if S.get("cls_results"):
            best_c = max(S["cls_results"], key=lambda k: S["cls_results"][k]["acc"])
            lines.append(f"   • Use {best_c} for classification tasks")
        lines.append("   • Run Predict New Data (Tab 12) for inference")
        lines.append("   • Export model package for deployment")
        lines.append("")
        lines.append("=" * 65)
        lines.append("   Report generated by ML Engine — Stage 1 + Stage 2")
        lines.append("=" * 65)

        S["final_report_text"] = "\n".join(lines)




 # ── Display ───────────────────────────────────────────────────────────
 # ─────────────────────────────────────────────────────────────────────────────
 # TAB 13 — DISPLAY + EXPORT  (replace your existing display/export block)
 # Same professional format as Tab 11
 # ─────────────────────────────────────────────────────────────────────────────

 # ── Display ───────────────────────────────────────────────────────────────────
if S.get("final_report_text"):
     st.text_area("📋 Final Report",
                  value=S["final_report_text"],
                  height=520,
                  key="final_report_area")

     st.markdown("---")
     section("📄 Export Final Report")
     ex1, ex2, ex3 = st.columns(3)

     # ── TXT ──────────────────────────────────────────────────────────────────
     with ex1:
         st.markdown("**Plain Text**")
         st.download_button(
             "📥 Download TXT",
             data=S["final_report_text"],
             file_name="Final_Insights_Report.txt",
             mime="text/plain",
             key="dl_final_txt",
             use_container_width=True
         )

     # ── PDF ──────────────────────────────────────────────────────────────────
     with ex2:
         st.markdown("**PDF Report**")
         if st.button("📄 Generate PDF", key="btn_final_pdf",
                      use_container_width=True):
             if not REPORTLAB_OK:
                 st.error("Install reportlab: `pip install reportlab`")
             else:
                 buf_pdf = io.BytesIO()
                 doc_pdf = SimpleDocTemplate(
                     buf_pdf, pagesize=A4,
                     topMargin=.7*inch, bottomMargin=.7*inch,
                     leftMargin=.85*inch, rightMargin=.85*inch)
                 styles = getSampleStyleSheet()
                 story  = []

                 # ── Styles ────────────────────────────────────────────────
                 title_style = ParagraphStyle(
                     "FinalTitle",
                     parent=styles["Title"], fontSize=20,
                     textColor=colors.HexColor("#1E3A8A"), spaceAfter=6)
                 h2_style = ParagraphStyle(
                     "FinalH2",
                     parent=styles["Heading2"], fontSize=13,
                     textColor=colors.HexColor("#1E3A8A"), spaceBefore=12)
                 h3_style = ParagraphStyle(
                     "FinalH3",
                     parent=styles["Heading3"], fontSize=11,
                     textColor=colors.HexColor("#059669"), spaceBefore=8)
                 body_style = ParagraphStyle(
                     "FinalBody",
                     parent=styles["Normal"], fontSize=9,
                     leading=14, spaceAfter=4)

                 # ── Title block ───────────────────────────────────────────
                 story.append(Paragraph(
                     "Final Insights & Recommendations", title_style))
                 story.append(Paragraph(
                     f"ML Engine — End-to-End · "
                     f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M}",
                     styles["Normal"]))
                 story.append(HRFlowable(
                     width="100%", thickness=2,
                     color=colors.HexColor("#1E3A8A")))
                 story.append(Spacer(1, .2*inch))

                 # ── Section 1: Dataset Summary ────────────────────────────
                 story.append(Paragraph("1. Dataset Summary", h2_style))
                 ds_data = [["Item", "Value"]]
                 if S.get("file_name"):
                     ds_data.append(["File", S["file_name"]])
                 if S.get("df_work") is not None:
                     df_w = S["df_work"]
                     ds_data.append(["Shape",
                         f"{df_w.shape[0]:,} rows × {df_w.shape[1]} cols"])
                 if S.get("target_col"):
                     ds_data.append(["Target", S["target_col"]])
                 if S.get("important_vars"):
                     ds_data.append(["Features selected",
                         str(len(S["important_vars"]))])
                 ds_tbl = Table(ds_data, colWidths=[2*inch, 4*inch])
                 ds_tbl.setStyle(TableStyle([
                     ("BACKGROUND", (0,0), (-1,0),
                      colors.HexColor("#1E3A8A")),
                     ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                     ("FONTSIZE",   (0,0), (-1,-1), 9),
                     ("GRID",       (0,0), (-1,-1), .3,
                      colors.HexColor("#CBD5E1")),
                     ("ROWBACKGROUNDS", (0,1), (-1,-1),
                      [colors.white, colors.HexColor("#F1F5F9")]),
                 ]))
                 story.append(ds_tbl)
                 story.append(Spacer(1, .15*inch))

                 # ── Section 2: EDA Insights ───────────────────────────────
                 story.append(Paragraph("2. EDA Insights (Stage 1)", h2_style))
                 stage1_text = S.get("insights_text", "")
                 if stage1_text:
                     for line in stage1_text.split("\n"):
                         stripped = line.strip()
                         if stripped and not stripped.startswith("=") \
                                 and "ML Engine" not in stripped \
                                 and "Report generated" not in stripped:
                             story.append(Paragraph(stripped, body_style))
                 else:
                     story.append(Paragraph(
                         "⚠️ Run Tab 8 in Stage 1 first to include EDA insights.",
                         body_style))
                 story.append(Spacer(1, .15*inch))

                 # ── Section 3: Regression Results ─────────────────────────
                 story.append(Paragraph(
                     "3. Regression Model Results (Stage 2)", h2_style))
                 if S.get("reg_results"):
                     r_data = [["Model", "R²", "MAE", "RMSE", "CV R²"]]
                     sorted_reg = sorted(
                         S["reg_results"].items(),
                         key=lambda x: x[1]["r2"], reverse=True)
                     for name, r in sorted_reg:
                         r_data.append([
                             name,
                             f"{r['r2']:.4f}",
                             f"{r['mae']:,.0f}",
                             f"{r['rmse']:,.0f}",
                             f"{r['cv_mean']:.4f}±{r['cv_std']:.4f}"])
                     r_tbl = Table(r_data,
                                   colWidths=[1.8*inch,.8*inch,
                                              .9*inch,.9*inch,1.3*inch])
                     r_tbl.setStyle(TableStyle([
                         ("BACKGROUND", (0,0), (-1,0),
                          colors.HexColor("#1E3A8A")),
                         ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                         ("FONTSIZE",   (0,0), (-1,-1), 9),
                         ("GRID",       (0,0), (-1,-1), .3,
                          colors.HexColor("#CBD5E1")),
                         ("ROWBACKGROUNDS", (0,1), (-1,-1),
                          [colors.white, colors.HexColor("#F1F5F9")]),
                         ("ALIGN",      (1,1), (-1,-1), "CENTER"),
                     ]))
                     story.append(r_tbl)
                     best_r = sorted_reg[0][0]
                     story.append(Spacer(1, .08*inch))
                     story.append(Paragraph(
                         f"🏆 Best Regression Model: <b>{best_r}</b>  "
                         f"R² = {S['reg_results'][best_r]['r2']:.4f}",
                         body_style))
                 else:
                     story.append(Paragraph(
                         "⚠️ No regression results. Train models in Tab 9.",
                         body_style))
                 story.append(Spacer(1, .15*inch))

                 # ── Section 4: Classification Results ─────────────────────
                 story.append(Paragraph(
                     "4. Classification Model Results (Stage 2)", h2_style))
                 if S.get("cls_results"):
                     c_data = [["Model", "Accuracy", "F1",
                                "Precision", "Recall", "CV Acc"]]
                     sorted_cls = sorted(
                         S["cls_results"].items(),
                         key=lambda x: x[1]["acc"], reverse=True)
                     for name, r in sorted_cls:
                         c_data.append([
                             name,
                             f"{r['acc']*100:.2f}%",
                             f"{r['f1']:.4f}",
                             f"{r['precision']:.4f}",
                             f"{r['recall']:.4f}",
                             f"{r['cv_mean']*100:.1f}%"])
                     c_tbl = Table(c_data,
                                   colWidths=[1.7*inch,.85*inch,.65*inch,
                                              .75*inch,.65*inch,.85*inch])
                     c_tbl.setStyle(TableStyle([
                         ("BACKGROUND", (0,0), (-1,0),
                          colors.HexColor("#059669")),
                         ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                         ("FONTSIZE",   (0,0), (-1,-1), 9),
                         ("GRID",       (0,0), (-1,-1), .3,
                          colors.HexColor("#CBD5E1")),
                         ("ROWBACKGROUNDS", (0,1), (-1,-1),
                          [colors.white, colors.HexColor("#ECFDF5")]),
                         ("ALIGN",      (1,1), (-1,-1), "CENTER"),
                     ]))
                     story.append(c_tbl)
                     best_c = sorted_cls[0][0]
                     story.append(Spacer(1, .08*inch))
                     story.append(Paragraph(
                         f"🏆 Best Classification Model: <b>{best_c}</b>  "
                         f"Accuracy = {S['cls_results'][best_c]['acc']*100:.2f}%",
                         body_style))
                 else:
                     story.append(Paragraph(
                         "⚠️ No classification results. Train models in Tab 10.",
                         body_style))
                 story.append(Spacer(1, .15*inch))

                 # ── Section 5: Final Recommendations ──────────────────────
                 story.append(Paragraph(
                     "5. Final Recommendations", h2_style))
                 recs = [
                     "• Verify missing value imputation (Tab 6)",
                     "• Check VIF multicollinearity (Tab 7)",
                     "• Review outlier treatment (Tab 3)",
                 ]
                 if S.get("reg_results"):
                     best_r = max(S["reg_results"],
                                  key=lambda k: S["reg_results"][k]["r2"])
                     recs.append(f"• Use {best_r} for regression tasks")
                 if S.get("cls_results"):
                     best_c = max(S["cls_results"],
                                  key=lambda k: S["cls_results"][k]["acc"])
                     recs.append(f"• Use {best_c} for classification tasks")
                 recs.append("• Run Predict New Data (Tab 12) for inference")
                 for rec in recs:
                     story.append(Paragraph(rec, body_style))

                 story.append(Spacer(1, .2*inch))
                 story.append(HRFlowable(
                     width="100%", thickness=1,
                     color=colors.HexColor("#CBD5E1")))
                 story.append(Paragraph(
                     "Report generated by ML Engine — Stage 1 + Stage 2 · Mohamed",
                     styles["Normal"]))

                 doc_pdf.build(story)
                 buf_pdf.seek(0)
                 st.download_button(
                     "⬇️ Download PDF",
                     buf_pdf.getvalue(),
                     file_name="Final_Insights_Report.pdf",
                     mime="application/pdf",
                     key="dl_final_pdf")

     # ── Word ─────────────────────────────────────────────────────────────────
     with ex3:
         st.markdown("**Word Report (.docx)**")
         if st.button("📝 Generate Word", key="btn_final_word",
                      use_container_width=True):
             if not DOCX_OK:
                 st.error("Install python-docx: `pip install python-docx`")
             else:
                 wd_doc = Document()
                 wd_doc.core_properties.author = "ML Engine — Mohamed"

                 # ── Helper ────────────────────────────────────────────────
                 def add_table_13(doc, data_rows, headers):
                     tbl = doc.add_table(
                         rows=1 + len(data_rows), cols=len(headers))
                     tbl.style = "Light Shading Accent 1"
                     for i, h_txt in enumerate(headers):
                         cell = tbl.rows[0].cells[i]
                         cell.text = h_txt
                         cell.paragraphs[0].runs[0].bold = True
                     for r_i, row in enumerate(data_rows, 1):
                         for c_i, val in enumerate(row):
                             tbl.rows[r_i].cells[c_i].text = str(val)

                 # ── Title ─────────────────────────────────────────────────
                 h0 = wd_doc.add_heading(
                     "Final Insights & Recommendations", 0)
                 h0.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                 wd_doc.add_paragraph(
                     f"ML Engine — End-to-End  |  "
                     f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M}")
                 wd_doc.add_paragraph()

                 # ── Dataset Summary ───────────────────────────────────────
                 wd_doc.add_heading("1. Dataset Summary", 1)
                 ds_rows = []
                 if S.get("file_name"):
                     ds_rows.append(["File", S["file_name"]])
                 if S.get("df_work") is not None:
                     df_w = S["df_work"]
                     ds_rows.append(["Shape",
                         f"{df_w.shape[0]:,} rows × {df_w.shape[1]} cols"])
                 if S.get("target_col"):
                     ds_rows.append(["Target", S["target_col"]])
                 if S.get("important_vars"):
                     ds_rows.append(["Features selected",
                         str(len(S["important_vars"]))])
                 if ds_rows:
                     add_table_13(wd_doc, ds_rows, ["Item", "Value"])
                 wd_doc.add_paragraph()

                 # ── EDA Insights ──────────────────────────────────────────
                 wd_doc.add_heading("2. EDA Insights (Stage 1)", 1)
                 stage1_text = S.get("insights_text", "")
                 if stage1_text:
                     for line in stage1_text.split("\n"):
                         stripped = line.strip()
                         if stripped and not stripped.startswith("=") \
                                 and "ML Engine" not in stripped:
                             wd_doc.add_paragraph(stripped)
                 else:
                     wd_doc.add_paragraph(
                         "⚠️ Run Tab 8 in Stage 1 to include EDA insights.")
                 wd_doc.add_paragraph()

                 # ── Regression Results ────────────────────────────────────
                 wd_doc.add_heading(
                     "3. Regression Model Results (Stage 2)", 1)
                 if S.get("reg_results"):
                     sorted_reg = sorted(
                         S["reg_results"].items(),
                         key=lambda x: x[1]["r2"], reverse=True)
                     reg_rows = [[n,
                                  f"{r['r2']:.4f}",
                                  f"{r['mae']:,.0f}",
                                  f"{r['rmse']:,.0f}",
                                  f"{r['cv_mean']:.4f}"]
                                 for n, r in sorted_reg]
                     add_table_13(wd_doc, reg_rows,
                                  ["Model","R²","MAE","RMSE","CV R²"])
                     best_r = sorted_reg[0][0]
                     wd_doc.add_paragraph(
                         f"🏆 Best: {best_r}  "
                         f"R² = {S['reg_results'][best_r]['r2']:.4f}")
                 else:
                     wd_doc.add_paragraph(
                         "⚠️ No regression results yet.")
                 wd_doc.add_paragraph()

                 # ── Classification Results ────────────────────────────────
                 wd_doc.add_heading(
                     "4. Classification Model Results (Stage 2)", 1)
                 if S.get("cls_results"):
                     sorted_cls = sorted(
                         S["cls_results"].items(),
                         key=lambda x: x[1]["acc"], reverse=True)
                     cls_rows = [[n,
                                  f"{r['acc']*100:.2f}%",
                                  f"{r['f1']:.4f}",
                                  f"{r['precision']:.4f}",
                                  f"{r['recall']:.4f}"]
                                 for n, r in sorted_cls]
                     add_table_13(wd_doc, cls_rows,
                                  ["Model","Accuracy","F1",
                                   "Precision","Recall"])
                     best_c = sorted_cls[0][0]
                     wd_doc.add_paragraph(
                         f"🏆 Best: {best_c}  "
                         f"Accuracy = "
                         f"{S['cls_results'][best_c]['acc']*100:.2f}%")
                 else:
                     wd_doc.add_paragraph(
                         "⚠️ No classification results yet.")
                 wd_doc.add_paragraph()

                 # ── Final Recommendations ─────────────────────────────────
                 wd_doc.add_heading("5. Final Recommendations", 1)
                 recs = [
                     "Verify missing value imputation (Tab 6)",
                     "Check VIF multicollinearity (Tab 7)",
                     "Review outlier treatment (Tab 3)",
                 ]
                 if S.get("reg_results"):
                     best_r = max(S["reg_results"],
                                  key=lambda k: S["reg_results"][k]["r2"])
                     recs.append(f"Use {best_r} for regression tasks")
                 if S.get("cls_results"):
                     best_c = max(S["cls_results"],
                                  key=lambda k: S["cls_results"][k]["acc"])
                     recs.append(f"Use {best_c} for classification tasks")
                 recs.append("Run Predict New Data (Tab 12) for inference")
                 for rec in recs:
                     wd_doc.add_paragraph(f"• {rec}")

                 wd_doc.add_paragraph()
                 wd_doc.add_paragraph(
                     "Report generated by ML Engine — Stage 1 + Stage 2 · Mohamed")

                 buf_wd = io.BytesIO()
                 wd_doc.save(buf_wd)
                 buf_wd.seek(0)
                 st.download_button(
                     "⬇️ Download Word",
                     buf_wd.getvalue(),
                     file_name="Final_Insights_Report.docx",
                     mime="application/vnd.openxmlformats-officedocument"
                          ".wordprocessingml.document",
                     key="dl_final_docx")

else:
     st.info("👆 Click **Generate Final Report** to combine "
             "Stage 1 + Stage 2 results.")


# ─────────────────────────────────────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown(
    f'<div style="text-align:center;color:#94A3B8;font-size:12px;">'
    f'ML Models Engine v2.0 · Stage 2 · '
    f'{datetime.datetime.now():%Y-%m-%d}'
    f'</div>',
    unsafe_allow_html=True)